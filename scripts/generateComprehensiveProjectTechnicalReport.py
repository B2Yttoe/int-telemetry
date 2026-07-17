# -*- coding: utf-8 -*-
"""Generate the comprehensive INT-Telemetry technical and experiment report.

The report is intentionally generated from source-controlled text and current
project artifacts so that the final Word deliverable can be reproduced after
future experiments are refreshed.
"""

from __future__ import annotations

import csv
import json
import math
import os
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "project-docs" / "INT-Telemetry_项目改进与实验完整技术报告.docx"
ASSET_DIR = ROOT / ".tmp" / "comprehensive-project-report-assets"

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_WIDTH = Inches(6.5)

BLUE = "215A8E"
DARK_BLUE = "173A5E"
CYAN = "158A9C"
GREEN = "2D7D5A"
ORANGE = "C26A2D"
RED = "B64242"
LIGHT_BLUE = "EAF3FA"
LIGHT_CYAN = "E9F6F7"
LIGHT_GREEN = "EAF5EF"
LIGHT_ORANGE = "FBF1E8"
LIGHT_RED = "FBECEC"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D7DEE5"
DARK_GRAY = "36424E"

FONT_CJK = "Microsoft YaHei"
FONT_BODY = "Microsoft YaHei"
FONT_CODE = "Consolas"
FONT_MATH = "Cambria Math"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_header_footer_font(paragraph, size=8.5, color="6A7784") -> None:
    for run in paragraph.runs:
        run.font.name = FONT_CJK
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    paragraph.add_run(" 页")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("222A31")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT_CJK
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(12)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_CJK
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_GRAY)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 13, 6),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
        ("Heading 4", 10.5, DARK_BLUE, 7, 3),
    ):
        style = styles[name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    if "Figure Caption" not in styles:
        figure_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        figure_style = styles["Figure Caption"]
    figure_style.font.name = FONT_CJK
    figure_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    figure_style.font.size = Pt(9)
    figure_style.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    figure_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_style.paragraph_format.space_before = Pt(3)
    figure_style.paragraph_format.space_after = Pt(9)
    figure_style.paragraph_format.keep_with_next = False

    if "Small Note" not in styles:
        note_style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note_style = styles["Small Note"]
    note_style.font.name = FONT_CJK
    note_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    note_style.font.size = Pt(9)
    note_style.font.color.rgb = RGBColor.from_string("5A6670")
    note_style.paragraph_format.space_after = Pt(4)
    note_style.paragraph_format.line_spacing = 1.15

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = FONT_CODE
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string("23313D")
    code_style.paragraph_format.left_indent = Inches(0.15)
    code_style.paragraph_format.right_indent = Inches(0.15)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run("INT-Telemetry 项目改进与实验完整技术报告")
    set_repeat_header_footer_font(header)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    set_repeat_header_footer_font(footer)

    props = doc.core_properties
    props.title = "INT-Telemetry 项目改进与实验完整技术报告"
    props.subject = "LEO 卫星网络仿真、INT/OAM 全网感知、LEO-INT-MC 改进与实验"
    props.author = "INT-Telemetry Project"
    props.keywords = "LEO, Walker, TLE, SGP4, INT, OAM, INT-MC, telemetry, matrix completion"
    props.comments = "由当前仓库源代码、正式实验报告和验收结果生成。"
    return doc


def add_heading(doc: Document, text: str, level=1, page_break=False):
    if page_break and len(doc.paragraphs) > 0:
        doc.add_page_break()
    paragraph = doc.add_heading(text, level=level)
    return paragraph


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, style=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for p_style in abstract.findall(".//" + qn("w:pStyle")):
            if p_style.get(qn("w:val")) == "ListNumber":
                abstract_id = abstract.get(qn("w:abstractNumId"))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        for abstract in numbering.findall(qn("w:abstractNum")):
            num_fmt = abstract.find(".//" + qn("w:numFmt"))
            if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
                abstract_id = abstract.get(qn("w:abstractNumId"))
                break
    if abstract_id is None:
        for item in items:
            doc.add_paragraph(item, style="List Number")
        return

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_ref)


def add_callout(doc: Document, title: str, text: str, kind="info") -> None:
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_GREEN, GREEN),
        "warning": (LIGHT_ORANGE, ORANGE),
        "danger": (LIGHT_RED, RED),
        "neutral": (LIGHT_GRAY, DARK_GRAY),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    table.columns[0].width = Inches(0.08)
    table.columns[1].width = Inches(6.25)
    left, body = table.rows[0].cells
    set_cell_width(left, 120)
    set_cell_width(body, 9000)
    set_cell_shading(left, accent)
    set_cell_shading(body, fill)
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(body, 110, 150, 110, 150)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(accent)
    p2 = body.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths: Sequence[int] | None = None,
    font_size=8.8,
    first_col_bold=False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    header = table.rows[0]
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "DCEAF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(text))
        run.bold = True
        run.font.name = FONT_CJK
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFB")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("" if value is None else str(value))
            run.font.name = FONT_CJK
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
            run.font.size = Pt(font_size)
            if first_col_bold and col_idx == 0:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    set_table_borders(table, color="DCE2E8", size="3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F9")
    set_cell_margins(cell, 100, 140, 100, 140)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in text.strip("\n").splitlines():
        p = cell.add_paragraph(style="Code Block")
        p.add_run(line)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc: Document, path: Path, caption: str, width=CONTENT_WIDTH, alt=None) -> bool:
    if not path.exists():
        add_callout(doc, "图像缺失", f"生成报告时未找到：{path.relative_to(ROOT)}", "warning")
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=width)
    set_alt_text(inline, caption, alt or caption)
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.paragraph_format.keep_with_next = False
    return True


def formula_image(latex: str, key: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / f"formula-{key}.png"
    if path.exists():
        return path
    fig = plt.figure(figsize=(8.2, 0.62), dpi=220, facecolor="white")
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=17, color="#173A5E")
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.09, facecolor="white")
    plt.close(fig)
    return path


def add_formula(doc: Document, latex: str, key: str, explanation: str | None = None) -> None:
    path = formula_image(latex, key)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = explanation is not None
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(5.9))
    set_alt_text(inline, f"公式 {key}", latex)
    if explanation:
        p = doc.add_paragraph(explanation, style="Small Note")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def read_experiment2_rows():
    path = ROOT / "reports" / "experiment2-int-mc-oracle-free-replay" / "experiment2-int-mc-enhancement-comparison.csv"
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_experiment3_rows():
    path = ROOT / "reports" / "experiment3-statistical-strong-baselines" / "experiment3-statistical-aggregate.csv"
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_experiment3_evidence():
    path = ROOT / "reports" / "experiment3-statistical-strong-baselines" / "experiment3-statistical-paired-evidence.csv"
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_experiment3_summary():
    path = ROOT / "reports" / "experiment3-statistical-strong-baselines" / "experiment3-statistical-summary.json"
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def percent_change(before: float, after: float) -> float:
    if before == 0:
        return 0.0
    return (after - before) / before * 100.0


def fmt(value, digits=4):
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def configure_matplotlib() -> FontProperties:
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if font_path.exists():
        prop = FontProperties(fname=str(font_path))
        plt.rcParams["font.family"] = prop.get_name()
    else:
        prop = FontProperties(family="sans-serif")
    plt.rcParams["axes.unicode_minus"] = False
    return prop


def create_architecture_figure() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "project-architecture.png"
    prop = configure_matplotlib()
    fig, ax = plt.subplots(figsize=(12, 6.3), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.3)
    ax.axis("off")
    layers = [
        (0.5, 4.65, 11.0, 1.1, "第一阶段：LEO 卫星网络真值黑盒", "TLE/SGP4 · Walker · 动态拓扑 · 链路预算 · 业务/路由 · 节点资源/能耗", "#EAF3FA", "#215A8E"),
        (0.5, 2.95, 11.0, 1.1, "第二阶段：INT 主动遥测与 Ground OAM", "probe 规划 · selective metadata · 报告回传 · 非全知重构 · INT-MC/强基线", "#E9F6F7", "#158A9C"),
        (0.5, 1.25, 11.0, 1.1, "第三阶段：小规模逐包交叉验证", "ns-3：INT header/MTU · 队列/丢包 · 报告交付 · AoI · 聚合模型趋势一致性", "#EAF5EF", "#2D7D5A"),
    ]
    for x, y, w, h, title, subtitle, fill, edge in layers:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.08", linewidth=1.8, edgecolor=edge, facecolor=fill)
        ax.add_patch(box)
        ax.text(x + 0.3, y + 0.72, title, fontsize=15, fontweight="bold", color=edge, fontproperties=prop)
        ax.text(x + 0.3, y + 0.31, subtitle, fontsize=10.5, color="#36424E", fontproperties=prop)
    for start, end in ((4.65, 4.05), (2.95, 2.35)):
        arrow = FancyArrowPatch((6, start), (6, end), arrowstyle="-|>", mutation_scale=16, linewidth=1.6, color="#667784")
        ax.add_patch(arrow)
    ax.text(6, 0.58, "实验 1–14：真实性、开销、误差、动态性、消融、因果合法性、系统级与外部盲测证据", ha="center", fontsize=11.5, color="#173A5E", fontproperties=prop)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def create_evidence_score_figure() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "current-evidence-score.png"
    prop = configure_matplotlib()
    labels = ["研究问题", "核心算法创新", "系统实现", "实验说服力", "外部真实性", "INFOCOM 就绪度"]
    scores = [8.5, 7.8, 8.5, 8.0, 6.3, 7.6]
    colors = ["#215A8E", "#158A9C", "#2D7D5A", "#6B7F99", "#C26A2D", "#785A9C"]
    fig, ax = plt.subplots(figsize=(10.5, 4.8), dpi=180)
    bars = ax.barh(labels[::-1], scores[::-1], color=colors[::-1], height=0.58)
    ax.set_xlim(0, 10)
    ax.set_xlabel("审计评分（10 分制）", fontproperties=prop)
    ax.grid(axis="x", linestyle="--", alpha=0.25)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="y", length=0)
    for tick in ax.get_yticklabels() + ax.get_xticklabels():
        tick.set_fontproperties(prop)
    for bar, value in zip(bars, scores[::-1]):
        ax.text(value + 0.12, bar.get_y() + bar.get_height() / 2, f"{value:.1f}", va="center", fontsize=10, color="#36424E", fontproperties=prop)
    ax.set_title("当前项目证据成熟度：技术链完整，外部真实性仍是主要边界", fontsize=14, color="#173A5E", fontproperties=prop)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def create_experiment2_delta_figure(rows) -> Path | None:
    if not rows:
        return None
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "experiment2-enhancement-deltas.png"
    prop = configure_matplotlib()
    metrics = [
        ("cpu_mae", "CPU MAE"),
        ("queue_depth_mae", "队列 MAE"),
        ("energy_percent_mae", "电量 MAE"),
        ("link_utilization_mae", "链路利用率 MAE"),
        ("utilization_inferred_mae", "补全利用率 MAE"),
    ]
    constellation_names = ["Iridium 66", "Telesat 351", "Starlink 1584"]
    data = []
    for name in constellation_names:
        before = next((r for r in rows if r["constellation_short_label"] == name and r["version"] == "before"), None)
        after = next((r for r in rows if r["constellation_short_label"] == name and r["version"] == "after"), None)
        if not before or not after:
            continue
        data.append([percent_change(float(before[key]), float(after[key])) for key, _ in metrics])
    if len(data) != 3:
        return None
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.7), dpi=180, sharey=True)
    for ax, name, values in zip(axes, constellation_names, data):
        colors = ["#2D7D5A" if value <= 0 else "#B64242" for value in values]
        bars = ax.barh([label for _, label in metrics][::-1], values[::-1], color=colors[::-1], height=0.58)
        ax.axvline(0, color="#36424E", linewidth=0.8)
        ax.set_title(name, fontproperties=prop, color="#173A5E")
        ax.grid(axis="x", linestyle="--", alpha=0.2)
        for tick in ax.get_yticklabels() + ax.get_xticklabels():
            tick.set_fontproperties(prop)
        for bar, value in zip(bars, values[::-1]):
            align = "left" if value >= 0 else "right"
            offset = 1.0 if value >= 0 else -1.0
            ax.text(value + offset, bar.get_y() + bar.get_height() / 2, f"{value:+.1f}%", va="center", ha=align, fontsize=8.5, fontproperties=prop)
    axes[0].set_ylabel("误差指标", fontproperties=prop)
    fig.suptitle("实验 2：增强方案相对原生 INT-MC 的误差变化（负值表示改善）", fontsize=14, color="#173A5E", fontproperties=prop)
    fig.tight_layout(rect=(0, 0, 1, 0.92))
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def create_experiment3_baseline_figure() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "experiment3-strong-baseline-summary.png"
    prop = configure_matplotlib()
    rows = read_experiment3_rows()
    evidence = read_experiment3_evidence()
    profiles = ["Telesat 351", "Starlink 1584"]
    metrics = ["CPU 利用率", "节点队列深度", "电量 SoC", "链路利用率"]
    backends = ["SoftImpute", "Kalman/RTS", "图邻居插值", "图时序正则", "ST-GNN", "CoSTCo"]
    significant = {
        (row.get("constellation_label"), row.get("metric_label"), row.get("backend_label"))
        for row in evidence
        if row.get("statistically_better_than_low_rank", "").lower() == "true"
    }
    by_key = {
        (row.get("constellation_label"), row.get("metric_label"), row.get("backend_label")): float(row.get("relative_mae_vs_low_rank", 0)) * 100
        for row in rows
    }
    fig, axes = plt.subplots(2, 1, figsize=(11.4, 8.2), dpi=180)
    norm = matplotlib.colors.TwoSlopeNorm(vmin=-60, vcenter=0, vmax=150)
    for ax, profile in zip(axes, profiles):
        values = [[by_key.get((profile, metric, backend), 0.0) for metric in metrics] for backend in backends]
        clipped = [[max(-60, min(150, value)) for value in row] for row in values]
        ax.imshow(clipped, cmap="RdYlGn_r", norm=norm, aspect="auto")
        ax.set_xticks(range(len(metrics)), metrics)
        ax.set_yticks(range(len(backends)), backends)
        ax.set_title(f"{profile}：相对原生低秩的 inferred-only MAE 变化", fontsize=12.5, color="#173A5E", fontproperties=prop, pad=10)
        for tick in ax.get_xticklabels() + ax.get_yticklabels():
            tick.set_fontproperties(prop)
            tick.set_fontsize(9)
        for y, backend in enumerate(backends):
            for x, metric in enumerate(metrics):
                value = values[y][x]
                marker = "*" if (profile, metric, backend) in significant else ""
                text_color = "white" if value <= -35 or value >= 90 else "#172033"
                ax.text(x, y, f"{value:+.1f}%{marker}", ha="center", va="center", fontsize=8.6, color=text_color, fontproperties=prop)
        for spine in ax.spines.values():
            spine.set_visible(False)
    fig.suptitle("实验 3：同观测、同字节、三种子、48 时间片强补全基线", fontsize=14, color="#173A5E", fontproperties=prop, y=0.99)
    fig.text(0.5, 0.012, "绿色/负值表示误差降低，红色/正值表示误差升高；* 表示配对移动块 bootstrap 的 95% CI 完全低于 0。颜色截断但单元格保留真实数值。", ha="center", fontsize=8.8, color="#5A6670", fontproperties=prop)
    fig.subplots_adjust(left=0.18, right=0.985, top=0.93, bottom=0.07, hspace=0.42)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def add_source_locations(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    add_table(doc, ["模块/文件", "关键位置", "职责"], rows, widths=[3500, 1500, 4360], font_size=8.2, first_col_bold=True)


def add_improvement(
    doc: Document,
    title: str,
    motivation: str,
    principle: Sequence[str],
    implementation: Sequence[str],
    effects: str,
    locations: Sequence[Sequence[str]],
    formulas: Sequence[tuple[str, str, str]] = (),
    boundary: str | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_callout(doc, "为什么提出", motivation, "info")
    add_heading(doc, "原理与底层逻辑", level=3)
    for paragraph in principle:
        add_paragraph(doc, paragraph)
    for latex, key, explanation in formulas:
        add_formula(doc, latex, key, explanation)
    add_heading(doc, "工程落实", level=3)
    add_bullets(doc, implementation)
    add_heading(doc, "关键代码落点", level=3)
    add_source_locations(doc, locations)
    add_callout(doc, "最终作用", effects, "success")
    if boundary:
        add_callout(doc, "证据边界", boundary, "warning")


def add_experiment(
    doc: Document,
    number: str,
    title: str,
    purpose: str,
    design: Sequence[str],
    metrics: Sequence[str],
    results: Sequence[str],
    conclusion: str,
    artifact: str,
    boundary: str | None = None,
    formula: tuple[str, str, str] | None = None,
) -> None:
    add_heading(doc, f"实验 {number}：{title}", level=2)
    add_callout(doc, "实验目的", purpose, "info")
    add_heading(doc, "实验设计", level=3)
    add_numbered(doc, design)
    add_heading(doc, "评价指标", level=3)
    add_bullets(doc, metrics)
    if formula:
        add_formula(doc, *formula)
    add_heading(doc, "结果与观察", level=3)
    add_bullets(doc, results)
    add_callout(doc, "实验结论", conclusion, "success")
    add_paragraph(doc, f"主要产物：{artifact}", style="Small Note")
    if boundary:
        add_callout(doc, "不能据此声称", boundary, "warning")


def build_report() -> Document:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    exp2_rows = read_experiment2_rows()
    exp3_rows = read_experiment3_rows()
    exp3_evidence = read_experiment3_evidence()
    exp3_summary = read_experiment3_summary()
    exp3_significant_count = sum(
        1 for row in exp3_evidence if row.get("statistically_better_than_low_rank", "").lower() == "true"
    )
    exp3_fairness_passed = bool(exp3_summary.get("fairness_audit", {}).get("all_cases_passed"))

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INT-TELEMETRY")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(CYAN)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("项目改进与实验完整技术报告")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("面向动态 LEO 卫星网络的高可信仿真、INT/OAM 全网感知与增强 INT-MC")
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    add_callout(
        doc,
        "报告定位",
        "本报告按论文技术文档的标准，系统记录项目从 Walker/TLE-SGP4 真值模型、节点与链路物理建模，到 INT 遥测、Ground OAM 重构、LEO-INT-MC 增强、强基线、消融、动态压力、拓扑复用、ns-3 逐包交叉验证和多源外部验证的完整演进。所有结论均区分“已证明”“条件成立”“尚未证明”，不以演示结果替代证据。",
        "neutral",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(45)
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_repeat_table_header(meta.rows[0])
    set_table_borders(meta, color="FFFFFF", size="0")
    cover_rows = [
        ("项目目录", str(ROOT)),
        ("报告日期", datetime.now().strftime("%Y-%m-%d")),
        ("覆盖阶段", "第一阶段真值模型 · 第二阶段 INT/OAM · 第三阶段逐包交叉验证"),
        ("实验范围", "实验 1–14（含 14B 前瞻性冻结快照协议）"),
        ("证据原则", "等预算、严格时间因果、无 oracle、外部验证不回调参数"),
    ]
    for idx, (key, value) in enumerate(cover_rows):
        left, right = meta.rows[idx].cells
        set_cell_width(left, 1900)
        set_cell_width(right, 7460)
        set_cell_margins(left, 70, 120, 70, 120)
        set_cell_margins(right, 70, 120, 70, 120)
        left.paragraphs[0].add_run(key).bold = True
        right.paragraphs[0].add_run(value)
        left.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_page_break()

    # Document control and abstract
    add_heading(doc, "文档控制与结论摘要", level=1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["文档性质", "研究级技术报告与实验日志；不是运营商设备规格书，也不是在轨部署证明。"],
            ["数据来源", "当前仓库源代码、配置、正式实验结果、验收脚本及冻结的外部数据快照。"],
            ["核心研究问题", "能否在动态、大规模 LEO 星座中，以严格遥测预算获得可信的全网节点/链路状态重构，并降低拓扑变化带来的规划成本。"],
            ["主要方法", "拓扑版本化、风险感知、单位字节信息增益、选择性 metadata、无 oracle OAM 闭环、结构/物理先验重构。"],
            ["当前成熟度", "研究问题 8.5/10；算法创新 7.5–8/10；系统实现 8.5/10；实验说服力约 8/10；外部真实性约 6–6.5/10。"],
        ],
        widths=[2100, 7260],
        font_size=9.1,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "一句话总评",
        "项目已经形成了“高可信聚合仿真真值黑盒 + 非全知 INT/OAM 感知 + 面向 LEO 的主动遥测与重构 + 小规模逐包交叉验证”的完整闭环，足以支撑论文初稿和后续严谨实验；但它仍不能声称复刻 Starlink 运营商内部网络，外部真实性与拓扑复用跨压力统计证据仍是主要边界。",
        "success",
    )
    add_heading(doc, "摘要", level=2)
    add_paragraph(doc, "INT-Telemetry 的目标不是构造一个只会画轨道的可视化演示，而是建立一个可向卫星网络投喂业务数据、按时间片产生节点和链路状态、再以非全知 INT 遥测方式恢复全网状态的研究平台。第一阶段负责生成真值：轨道由 Walker 或 TLE/GP/OMM 与 SGP4 传播，链路由距离、地球遮挡、极区、天线端口和链路预算共同决定，节点资源由业务、转发、排队、光照和能量方程驱动。第二阶段把第一阶段视为黑盒，仅通过 probe、逐跳 metadata、reporting path 和 Ground OAM 获取部分观测，再使用原生 INT-MC、增强 LEO-INT-MC 及多种强基线重构全网状态。第三阶段抽取小规模动态窗口，以 ns-3 检验 INT header、MTU、逐包排队、丢包、报告交付与 AoI 是否和聚合模型保持一致趋势。")
    add_paragraph(doc, "本项目的主要算法创新逐步收敛为：在预测拓扑和历史 OAM 部分观测下，以拓扑版本、状态不确定性、未来断链风险和硬字节预算为四个核心量，在 reuse、local repair、fresh 三类规划方式及 full、compact、forward-only 三类 metadata 动作之间进行统一决策；OAM 重构层再利用活动掩码、时间因果、轨道邻居、节点—链路耦合与电量物理先验完成状态恢复。实验表明，这一方法在中大型星座中能够以相近或条件更低开销改善多个重构指标，并在 Telesat 351 的 48 片窗口中证明 local repair 可减少确定性规划工作量；同时，负结果也表明增强方案并非在所有规模、所有指标和所有开销上普遍胜出。")
    add_figure(doc, create_architecture_figure(), "图 1  项目三阶段技术链与实验验证关系", width=Inches(6.35))

    add_heading(doc, "目录", level=1, page_break=True)
    toc_rows = [
        ["第 1 章", "项目目标、研究边界与技术路线"],
        ["第 2 章", "第一阶段：卫星网络真值模型的逐项改进"],
        ["第 3 章", "第二阶段：INT、Ground OAM 与非全知边界"],
        ["第 4 章", "原生 INT-MC 迁移到动态 LEO 的问题"],
        ["第 5 章", "增强 LEO-INT-MC 的算法与工程改进"],
        ["第 6 章", "重构后端、强基线与多指标探索"],
        ["第 7 章", "实验方法、评价指标和公平性原则"],
        ["第 8 章", "实验 1–14 的完整记录与结论"],
        ["第 9 章", "已证明、条件成立与尚未证明的结论"],
        ["第 10 章", "当前成熟度、论文立足点与后续工作"],
        ["附录", "符号、关键配置、代码索引与复现入口"],
    ]
    add_table(doc, ["章节", "内容"], toc_rows, widths=[1500, 7860], font_size=9.5, first_col_bold=True)
    add_paragraph(doc, "提示：本目录为稳定的章节导航；在 Word 中可通过“引用→目录”按标题样式生成带页码的自动目录。", style="Small Note")

    # Chapter 1
    add_heading(doc, "第 1 章  项目目标、研究边界与技术路线", level=1, page_break=True)
    add_heading(doc, "1.1 项目最终要解决什么问题", level=2)
    add_paragraph(doc, "项目围绕一个清晰的问题展开：动态 LEO 星座的拓扑在轨道运动、极区切换和星地窗口作用下持续变化，卫星节点和星间链路状态又会被业务流、转发负载、队列、能量与通信环境共同影响。在这样的网络中，若每个时间片全量采集全部节点和链路 metadata，遥测开销会很高；若只测很少对象，又会出现大面积 unknown。项目因此需要同时解决“可信真值如何生成”“部分观测如何获得”“全网状态如何重构”“重构误差和遥测开销如何权衡”四个问题。")
    add_formula(doc, r"\min_{\mathcal{P}_t,\hat{X}_t}\;\lambda_b C_{bytes}(\mathcal{P}_t)+\lambda_e C_{energy}(\mathcal{P}_t)+\lambda_p C_{plan}(\mathcal{P}_t)+\lambda_r\mathbb{E}[\mathcal{E}_t]+\lambda_a AoI_t", "1-1", "统一研究目标：在通信、能量和规划成本受控时，降低 OAM 重构误差与状态年龄。")
    add_heading(doc, "1.2 三阶段职责边界", level=2)
    add_table(
        doc,
        ["阶段", "可见信息", "核心职责", "不能替代的内容"],
        [
            ["第一阶段真值模型", "全知节点、链路、轨道、业务和路由状态", "生成动态卫星网络及每时间片真值；导出可复现实验数据", "不能证明真实运营商内部状态逐点一致"],
            ["第二阶段 INT/OAM", "仅有 probe 直接观测、报告回传和历史 OAM", "低开销采样、报告、重构、误差检验和数据下载", "过程页不能提前读取第一阶段隐藏真值"],
            ["第三阶段 ns-3", "逐包业务、probe、report、队列和 MTU", "检验聚合模型的包级可执行性与趋势", "不承担 1584 星长窗口全量算法搜索"],
        ],
        widths=[1700, 2350, 2800, 2510],
        font_size=8.4,
        first_col_bold=True,
    )
    add_heading(doc, "1.3 核心研究主张与边界", level=2)
    add_bullets(doc, [
        "可主张：项目能在公开轨道与研究级通信/能耗参数约束下，稳定生成物理一致、可解释、可复现的节点和链路状态。",
        "可主张：第二阶段在非全知、无 oracle、严格时间因果条件下完成 INT 采集、Ground OAM 重构和真值后验检验。",
        "可主张：拓扑版本化、风险感知、单位开销信息增益和硬预算可以统一为一个主动遥测规划问题。",
        "条件主张：在已验证的中大型星座、窗口与压力范围内，增强方案形成等预算或相近预算的重构质量优势。",
        "不可主张：CPU、电量、队列与 Starlink 运营商在轨内部真值逐点一致；真实 ISL 调度、路由策略和硬件参数并未公开。",
        "不可主张：当前所有机器学习补全后端均优于低秩补全，或拓扑 exact reuse 已在所有规模和压力下取得统计显著收益。",
    ])

    # Chapter 2: Stage 1 improvements
    add_heading(doc, "第 2 章  第一阶段：卫星网络真值模型的逐项改进", level=1, page_break=True)
    add_paragraph(doc, "第一阶段已经被封装为后续研究的真值黑盒。以下改进按实际演进顺序归纳，每项都回答为何需要、如何实现、产生什么作用，以及证据能够覆盖到哪里。")

    add_improvement(
        doc,
        "2.1 从示意节点图改为标准 Walker 轨道结构",
        "早期拓扑只满足节点均匀排列，不能保证相邻卫星在同一轨道面内沿轨相邻，也不能表达 Walker 相位因子。错误的 plane-slot 映射会导致轨内链路跨越地球或穿模，并把本应短时延的相邻链路变成长路径。",
        [
            "Walker 星座用轨道面数 P、每面卫星数 S、总星数 T=P×S 和相位因子 F 描述。轨道面升交点赤经在 180° 或 360° 范围内均匀划分，本项目的 Walker-Star 主配置采用 180° 平分，以避免把同一物理轨道平面的反向法向重复计数。",
            "每个槽位的平近点角在单轨内均匀分布，并加入跨轨道面的相位偏移。这样可从 plane_id 和 slot_id 唯一推回邻居，并保证轨内前向/后向链路是几何上的真实沿轨邻接。",
        ],
        [
            "在配置中显式维护 planes、satellitesPerPlane、phaseFactor、inclination 和 altitude。",
            "轨内链路固定连接 slot±1；跨面链路只在相邻轨道面和天线端口允许时建立。",
            "仪表盘增加轨道面显示开关，以轨道圆/轨道面帮助人工检查 plane-slot 映射。",
        ],
        "消除了轨内链路跨槽位和穿模的主要来源，使拓扑的轨内稳定性、轨间动态性和最大四链路约束能够在同一个几何模型中成立。",
        [["src/simulation/walker.ts", "Walker 生成与传播", "轨道面、槽位、位置与邻接关系"], ["src/config/walkerNetworkConfig.ts", "星座配置", "规模、相位、倾角、高度和链路阈值"]],
        formulas=[(r"\Omega_p=\Omega_0+\frac{\pi p}{P},\qquad M_{p,s}=M_0+\frac{2\pi s}{S}+\frac{2\pi Fp}{PS}", "2-1", "Walker-Star 轨道面与槽位相位关系。")],
        boundary="Walker 结构描述的是研究星座设计；除非使用冻结的公开 GP/OMM 并完成映射，否则不等同于某一时刻运营 Starlink 的真实 plane-slot 编排。",
    )

    add_improvement(
        doc,
        "2.2 接入 TLE/GP/OMM 与 SGP4 真实轨道模式",
        "只用理想圆轨道会忽略真实历元、偏心率、BSTAR、近地点幅角和实际相位差。为了让拓扑至少在公开轨道层面可验证，项目增加 TLE/GP/OMM 输入和 SGP4 传播模式，并保留可预测、可视的 Walker 模式作为对照。",
        [
            "每颗卫星保存 satellite_id、NORAD ID、COSPAR ID、名称、shell/plane/slot、TLE 两行、epoch、倾角、RAAN、偏心率、近地点幅角、平近点角、平均运动、BSTAR 和运行状态。",
            "TLE 与 SGP4 绑定使用：传播器在时间 t 输出 ECI 位置和速度，再转换为 ECEF、经纬度、高度、地面轨迹、轨道速度等时隙状态。",
            "为了兼顾论文复现和数据新鲜度，后续采用“单次新鲜轨道快照冻结 + 有界时间窗口外推”：输入 GP0 冻结，未来 GP1 仅用于盲测，不反向调参。",
        ],
        [
            "增加真实目录读取、TLE 解析、合成 TLE 和 SGP4 传播接口。",
            "通过 orbitMode 在 walker 与 tle-sgp4 之间切换，前端保持相同拓扑可视化接口。",
            "冻结快照时记录来源、获取时间、历元年龄、筛选规则和 SHA-256 哈希。",
        ],
        "轨道真实性从内部几何一致性提升到可由公开 GP/OMM 外部复核；时间片位置、速度、星下点和由此产生的链路几何都由传播结果驱动。",
        [["src/simulation/tle.ts", "TLE 解析/传播", "TLE 字段、SGP4 状态和坐标转换"], ["src/simulation/realTleCatalog.ts", "公开目录接入", "目录筛选、星座映射与快照"], ["scripts/experiments/runExperiment14B*", "实验 14B", "冻结、门禁、未来快照盲测"]],
        formulas=[(r"(\vec{r}_i(t),\vec{v}_i(t))=\mathrm{SGP4}(\mathrm{GP}_i,t)", "2-2", "公开轨道元素经 SGP4 传播为时间片轨道状态。"), (r"T=2\pi\sqrt{\frac{a^3}{\mu}},\qquad v\approx\sqrt{\frac{\mu}{a}}", "2-3", "圆轨道近似用于合理性检查，而不是替代 SGP4。")],
        boundary="公开 TLE/GP 可验证轨道位置，但不能公开推导运营商内部 ISL 是否启用、路由策略和链路调度。Starlink 1584 的 72×22 是标准主壳层研究规模，当前冻结样本的实际轨道群组参数仍以 manifest 为准。",
    )

    add_improvement(
        doc,
        "2.3 三种现实启发星座规模与统一接口",
        "仅用 8×8 或几十颗示意卫星无法研究大规模遥测开销和矩阵缺失比例。项目最终收敛为小、中、大三种公开可解释的规模，并让三者共享同一仿真、导出、INT 和实验接口。",
        [
            "小型采用 Iridium 66（6×11）作为低轨小规模对照；中型采用研究中常用的 Hypatia Telesat-1015 设计（27×13=351）；大型采用 Starlink 主壳层标准研究规模 72×22=1584。",
            "星座 profile 只替换轨道和设备/链路配置，不改变实验算法，避免为某一个规模写专门保护逻辑。",
        ],
        [
            "建立 constellation profile 验证脚本，检查节点数、plane-slot 唯一性、公开来源和轨道统计。",
            "正式实验只保留三种规模，历史 8×8、47×14 等开发规模归档，不再作为论文主结果。",
            "规模自适应预算使用统一函数，不能以 constellation name 直接分支调参。",
        ],
        "使实验能够观察小规模过采样、中规模局部复用和大规模采样稀释等不同现象，并为跨规模泛化提供统一比较基础。",
        [["src/config/constellationProfiles.ts", "星座 profile", "三种规模和来源元数据"], ["scripts/verifyConstellationProfiles.mjs", "配置验收", "数量、轨道和来源检查"], ["reports/experiment10-*", "多种子公平实验", "三规模同预算比较"]],
        boundary="Telesat 351 是公开研究设计而非在轨完整 TLE；Iridium/Starlink 的公开轨道可验证程度高于运营 ISL 状态可验证程度。",
    )

    add_improvement(
        doc,
        "2.4 动态链路生成：四端口、距离、遮挡与极区断链",
        "早期可视化曾出现进入极区后链路仍显示、断开状态与拓扑线条不一致、同一卫星重复建链和链路穿过地球等问题。需要把链路状态从绘图逻辑中剥离，由统一物理与端口约束生成，再让所有视图消费同一 active edge 集合。",
        [
            "每星抽象为 front、back、left、right 四个 ISL 端口，每个端口同时最多承载一条链路，因此节点度不超过 4，轨内前后链路保持稳定，轨间左右链路随时间变化。",
            "候选链路必须同时通过距离门限、地球视线遮挡、极区纬度、天线占用、指向/太阳规避等门禁。极区规则对两个端点同时检查，只要任一端进入限制区，轨间链路立即 down。",
            "二维快照、三维动态拓扑、链路状态表和导出 CSV 均读取同一时间片 link.status/active 字段，防止“表中断链但图上仍画线”。",
        ],
        [
            "轨内连接由 plane-slot 邻接固定生成；轨间连接由每时间片候选匹配生成。",
            "地球遮挡通过线段到地心最近距离与地球半径比较，不允许 ISL 穿过地球。",
            "链路去重使用无向规范化 link_id，并绑定唯一端口占用。",
        ],
        "动态拓扑由可解释门禁决定，极区进出会同步触发左右轨间链路断开/恢复，图、表、INT active mask 与导出状态保持一致。",
        [["src/simulation/walker.ts", "链路生成", "轨内/轨间候选与时间片状态"], ["src/simulation/antenna.ts", "端口分配", "四方向 ISL 和 SGL 天线占用"], ["src/components/*Topology*", "前端", "按当前 active link 绘制"]],
        formulas=[(r"A_{ij}(t)=I[d_{ij}\leq d_{max},\ LOS_{ij},\ |\phi_i|,|\phi_j|<\phi_{polar},\ ports\ free]", "2-4", "轨间链路活动掩码；任一门禁失败即为 down。")],
        boundary="最远通信距离和极区阈值是公开研究参数/配置，不代表 Starlink 未公开的实际激光终端调度阈值。",
    )

    add_improvement(
        doc,
        "2.5 独立天线模块与星地回传约束",
        "把“最多四链路”只写成拓扑规则无法解释为何不能继续建链，也无法研究星地回传窗口、天线占用和 metadata 报告最终能否送达 Ground OAM。",
        [
            "ISL 天线抽象为轨内前/后与轨间左/右四个方向端口；SGL 天线独立维护频段、增益、波束宽度、最远距离、发射功率、带宽、波束数、指向范围、转向速度和 idle/occupied/switching/fault 状态。",
            "星地链路由可见窗口、最低仰角、距离、指向和当前占用共同决定。INT sink 生成 report 后，还必须找到 reporting path 并在容量和窗口内下传，才能被 Ground OAM 接收。",
        ],
        [
            "建立 antenna.ts 独立模块，链路生成器只请求可用天线，不直接修改端口状态。",
            "将 SGL 报告容量、仰角门限和切换可用比例加入配置。",
            "第二阶段导出 delivered/undelivered report，Ground OAM 只消费已交付报告。",
        ],
        "天线从隐含度数限制变成可扩展设备资源，支持未来接入 HAPS、用户链路、多波束或故障状态；更重要的是把 INT 采集与星地回传约束连成完整闭环。",
        [["src/simulation/antenna.ts", "天线模型", "ISL/SGL 属性、状态与占用"], ["stage2-int/tools/reporting-path-planner.mjs", "报告路径", "sink 到地面 OAM 的回传规划"], ["stage2-int/tools/ground-oam-reconstructor.mjs", "OAM 输入", "仅接收成功交付报告"]],
        boundary="当前天线为研究级链路抽象，不模拟完整波束成形、射频前端非线性或硬件伺服过程。",
    )

    add_improvement(
        doc,
        "2.6 链路预算与高影响通信因素",
        "仅用距离随机生成容量和时延会削弱模型可信度；但完整硬件级射频/激光仿真又会让复杂度压过 Walker 拓扑与 INT 主题。因此项目保留高影响、低复杂度因素，暂缓精细物理层。",
        [
            "自由空间损耗由距离和载波频率决定；接收功率由发射功率、收发增益和总损耗求得；噪声由带宽、噪声谱密度和噪声系数计算；同频/邻频干扰聚合后形成 SINR；容量用带效率系数的 Shannon 上界估计。",
            "额外保留频率复用/邻道抑制、指向损耗、切换可用比例、多普勒及补偿残差、太阳规避角、SGL 雨衰和动态天气。精细极化、激光线宽、PLL、完整 FEC/MAC 等被明确暂缓。",
        ],
        [
            "linkBudget.ts 统一计算 FSPL、Pr、N、SNR/SINR 和容量，节点/链路导出保存中间物理量。",
            "spaceEnvironment.ts 生成太阳几何、天气和光照上下文。",
            "配置文件暴露设备参数，实验冻结配置哈希，避免运行中隐式变化。",
        ],
        "链路容量、质量、warning/down 状态和业务拥塞都由同一链路预算链条驱动；模型从“随机链路状态”提升为研究级物理一致状态，同时保持 1584 星长窗口实验可运行。",
        [["src/simulation/linkBudget.ts", "链路预算", "FSPL、接收功率、噪声、SINR、容量"], ["src/simulation/spaceEnvironment.ts", "空间环境", "太阳、天气和传播上下文"], ["src/config/walkerNetworkConfig.ts", "设备参数", "频率、功率、增益、带宽和阈值"]],
        formulas=[(r"L_{FS}=92.45+20\log_{10}d_{km}+20\log_{10}f_{GHz}", "2-5", "自由空间路径损耗，单位 dB。"), (r"P_r=P_t+G_t+G_r-L_{total},\quad N=-174+10\log_{10}B_{Hz}+NF", "2-6", "接收功率与热噪声功率的 dB 形式。"), (r"SINR=\frac{P_s}{N+\sum_k I_k},\qquad C=\eta B\log_2(1+SINR)", "2-7", "干扰聚合后的链路容量估计。")],
        boundary="容量是研究级有效吞吐上界/近似，不是完整 MCS、PER、FEC 编解码和 MAC 重传的逐包物理层结果。",
    )

    add_improvement(
        doc,
        "2.7 光照、太阳翼与电池状态方程",
        "早期电量由演示波动生成，可能出现不合理归零或光照面仍持续下降。为让电量能用于任务调度、INT 重构和后续机器学习，项目改为由太阳几何、太阳翼参数、负载功率和充放电效率逐时间片传播。",
        [
            "太阳正面时，太阳翼峰值功率由太阳常数、面积和效率计算；阴影面发电为零，半影按光照系数缩放。负载由基础、通信、计算和载荷功耗组成，并随任务与转发活动改变。",
            "电池维护 Wh 能量与 SoC。充电和放电采用不同效率，能量被裁剪到最小安全能量与额定容量之间；SoC 低于 20% 进入节能模式，而不是允许电量物理上变成负值。",
            "当前默认研究参数为太阳翼 2.0 m²、效率 0.28、电池 1200 Wh、基础/通信/计算/载荷 80/100/50/100 W、充放电效率 0.95。",
        ],
        [
            "在每个时间片先计算 sunlit/shadow/partial，再计算发电、负载和净功率。",
            "任务 CPU、转发活动、通信天线占用进入负载；空业务 operational 模式 CPU 接近 0，但基础平台仍耗电。",
            "第二阶段重构器复用相同可预测光照和静态负载上下文作为 energy physics prior。",
        ],
        "电量变化成为可解释的动态状态：光照面总体充电、阴影面放电、低电量触发节能，且可从上一时间片和功率收支复算，显著降低大规模星座电量补全 MAE。",
        [["src/simulation/walker.ts", "节点状态推进", "光照、负载与电量更新"], ["src/config/walkerNetworkConfig.ts", "电源参数", "太阳翼、电池、功耗与阈值"], ["stage2-int/tools/int-mc-energy-prior.mjs", "重构先验", "基于过去状态的因果能量传播"]],
        formulas=[(r"P_{gen}(t)=s(t)I_{sun}A_{sa}\eta_{sa}", "2-8", "光照系数 s(t)∈[0,1]。"), (r"E_{t+1}=\mathrm{clip}\!\left[E_t+\eta_{ch}[P_{gen}-P_{load}]_+\Delta t-\frac{[P_{load}-P_{gen}]_+}{\eta_{dis}}\Delta t\right]", "2-9", "离散时间电池状态方程。")],
        boundary="参数是论文级近似设备配置，不是运营商公开硬件真值；电池老化、热控、姿态伺服和精细太阳翼功率点跟踪未建模。",
    )

    add_improvement(
        doc,
        "2.8 节点资源、运行模式与业务驱动状态",
        "只用 seededWave 或随机数填充 CPU、内存和队列只能用于演示，不能支持“真实业务输入会让卫星产生合理响应”的研究目标。项目保留演示模式，同时新增 operational 模式，使状态由任务、转发、缓存、链路和能量公式驱动。",
        [
            "每个计算节点保存 node_id、node_type、CPU/GPU 容量、内存、存储、电量、CPU/内存/存储利用率、队列、健康状态、参与任务（源端/经过）及静态轨道信息。",
            "任务从源节点产生，经路由路径逐跳转发。CPU 包含源端计算、过境转发、排队管理和遥测处理；内存/存储与任务数据、队列和遥测缓存相关。空业务 operational 模式下任务 CPU 为 0，但基础平台和仿真控制开销不会被误报为业务 CPU。",
        ],
        [
            "仪表盘提供 demo/operational 切换；demo 保留可重复演示波动，operational 禁止随机填充业务状态。",
            "参与任务显示改为“源端/经过”，悬停可查看 task_id、源/目的、产生片、上一跳、下一跳和状态。",
            "静态轨道字段折叠，动态资源和任务状态保持首屏可读。",
        ],
        "节点状态成为业务和物理过程的响应变量，可用于 INT 节点 metadata、OAM 重构和机器学习数据集；同时保留演示模式便于 UI 展示，二者语义明确分离。",
        [["src/simulation/types.ts", "类型定义", "节点资源、任务参与和健康状态"], ["src/simulation/walker.ts", "状态方程", "CPU、内存、队列、能量与任务耦合"], ["src/components/*Node*", "仪表盘", "节点详情、任务悬浮和折叠字段"]],
        formulas=[(r"CPU_i(t)=\operatorname{clip}\{C_i^{src}+C_i^{transit}+C_i^{queue}+C_i^{telemetry},0,100\}", "2-10", "节点 CPU 利用率由源任务、转发、队列与遥测工作量共同形成。")],
        boundary="这些状态是物理/工作负载驱动的仿真潜变量，外部运营商没有公开逐星 CPU、内存和电量 trace，因此真实性只能通过方程、范围、敏感性和独立包级趋势验证。",
    )

    add_improvement(
        doc,
        "2.9 业务数据端口、长时业务与 Cloudflare Radar 校准",
        "如果只有短时、少任务或平滑合成流量，CPU、队列和链路利用率的动态范围不足，无法评价 INT 重构与后续预测。项目因此定义可上传业务数据端口，并扩展 48 时间片的长时、多批次、周期和热点迁移业务。",
        [
            "业务数据至少包含 task_id、generation_time/slice、source、destination、data_size、compute_demand、deadline/priority 和可选周期字段。输入先做节点、时间、大小、路径可达性和重复 ID 校验。",
            "Cloudflare Radar 的公开时序用于校准业务相对波动、峰谷和周期特征，而不是直接当作卫星运营商逐流 trace。正确流程是只用校准段拟合强度映射，测试段冻结后评估相关性，不允许根据测试结果回调参数。",
            "内置场景包括空业务、正常业务和高负载业务，外部 CSV 通过相同端口注入。",
        ],
        [
            "traffic.ts 负责解析、校验、场景生成、周期扩展和任务分配。",
            "导出任务产生、路由、完成、超时、丢弃、排队和端到端时延，便于后续监督学习。",
            "外部数据快照、校准模板和实验生成业务分目录保存，避免把派生数据误称为原始 Radar 数据。",
        ],
        "业务从静态演示输入提升为可复现的扰动源，48 个时间片中节点和链路状态有充分变化；Radar 能提升宏观业务节律可信度，但项目仍如实标注其语义差异。",
        [["src/simulation/traffic.ts", "业务模块", "导入、校验、场景与周期业务"], ["examples/datasets", "输入示例", "标准和长时业务 CSV"], ["external-data / reports/experiment14*", "外部校准", "Radar 冻结段与验证段"]],
        boundary="Cloudflare Radar 是互联网宏观流量指数，不是 Starlink 内部逐流 trace；0.9134 等相关系数若使用了校准段，只能证明拟合，不是独立泛化。",
    )

    add_improvement(
        doc,
        "2.10 拥塞感知最短路径、队列和任务时延修复",
        "早期任务 model_estimated_latency 的 P95 达到分钟级，原因包括 plane-slot 错配、把未发送流量在每跳重复累加为队列时延、以及对 dropped/partial 任务仍输出巨大 estimated_latency。该问题会污染真实性实验和后续 ML 标签。",
        [
            "路由由纯传播最短路径改为拥塞感知最短路径：边权包含传播时延、利用率惩罚、队列惩罚和链路状态门禁，避免把所有流量压到少数几条最短几何链路。",
            "队列延迟由链路实际队列字节与服务容量计算，并在路径上合理组合；未发送的剩余流量保留为 backlog，不在每一跳重复累计。",
            "完成任务输出完成时延；超时、丢弃和 partial 单独记录 timeout/drop/queue_backlog，不再用分钟级大值冒充有效 latency。",
        ],
        [
            "修正 TLE/Walker plane-slot 映射，确保轨内邻居沿轨相邻。",
            "设置 30 s 路由队列/任务超时上限，并导出分状态任务字段。",
            "保持用户侧 RTT 与内部任务完成时间口径分离：RIPE Atlas RTT 不再直接对比包含计算和等待的内部任务时延。",
        ],
        "修正了分钟级伪时延的主要漏洞，使 P50 用户侧延迟回到合理量级，并让 P95 高时延能够解释为拥塞/未完成尾部，而不是错误累加。",
        [["src/simulation/walker.ts", "路由/队列推进", "拥塞感知边权、服务与 backlog"], ["src/simulation/export.ts", "任务导出", "completed/partial/dropped/timeout 分栏"], ["src/config/walkerNetworkConfig.ts", "时延门限", "queue carryover、capacity 和 timeout"]],
        formulas=[(r"w_e(t)=d_e^{prop}+\alpha\,\phi(U_e(t))+\beta\,q_e(t)+\infty\cdot I[e\;down]", "2-11", "拥塞感知最短路径边权。"), (r"d_e^{queue}(t)\approx\frac{8Q_e(t)}{C_e(t)}", "2-12", "链路队列字节除以服务速率得到排队时延近似。")],
        boundary="聚合时间片队列不能完整复现微突发、逐包调度和重传尾时延，因此实验 13 又用 ns-3 做小规模交叉验证。",
    )

    add_improvement(
        doc,
        "2.11 全透明 3D/2D 仪表盘与时间片同步",
        "把所有控制、节点、链路、轨道和实验结果放在一个页面会造成信息拥挤；同时，运动视图与快照状态混用会让用户误解当前链路是否属于某个时间片。",
        [
            "3D 视图以地球为中心，轨道面均匀分布，卫星沿轨运动，整幅拓扑可旋转、缩放和点击；地球自转被取消，卫星相对地面的东西向运动通过轨道/ECEF 转换体现。",
            "未选择时间片时展示连续运动；选择 T00–T47 时冻结为快照，节点和链路严格同步到该片。二维拓扑额外展示同一时间点的平面化结构，方便分析路径和链路断开。",
            "UI 重排为导航、抽屉和独立模块窗口：控制台顶部可锁定/收起，时间片选择置顶悬停，其他模块按需拉起；内容字段不改变，只改善可读性。",
        ],
        [
            "节点、链路和轨道面均有独立显示开关。",
            "点击节点显示资源、轨道、参与任务、天线和链路状态。",
            "第一阶段真值页和第二阶段遥测过程页使用不同信息权限与视觉语义。",
        ],
        "仪表盘既能呈现动态轨道直觉，也能作为每时间片真值审计工具；控制台和模块不再挤压单一页面，实验结果也能独立展示。",
        [["src/App.tsx", "页面与导航", "第一阶段/遥测/实验入口"], ["src/components", "可视化组件", "3D、2D、抽屉、详情和筛选"], ["src/styles", "界面样式", "中文布局和响应式约束"]],
        boundary="3D 画面用于拓扑解释与人工审计，不是高精度轨道分析软件；真实位置精度由传播和数据快照决定，而不是渲染精度。",
    )
    add_figure(doc, ROOT / "project-docs" / "assets" / "stage1-dashboard-overview.png", "图 2  第一阶段仪表盘：3D 星座、状态概览与时间片控制", width=Inches(6.2))
    add_figure(doc, ROOT / "project-docs" / "assets" / "ui-redesign-topology.png", "图 3  重新排版后的拓扑与模块化导航界面", width=Inches(6.2))

    add_improvement(
        doc,
        "2.12 数据导出、配置冻结与第一阶段黑盒封装",
        "后续 INT、矩阵补全和机器学习需要稳定、可追溯的输入输出契约。若每次修改都直接读取前端内存对象，实验无法复现，也容易产生真值泄漏。",
        [
            "第一阶段输入包括星座 profile、轨道快照、时间范围、业务数据、链路/天线/能耗参数和随机种子；输出按时间片导出 nodes、links、routes、tasks、metadata 和配置/哈希。",
            "第二阶段只能通过显式端口读取拓扑可达性、probe 执行和已交付 metadata；真值 CSV 只在实验后评估器中使用。",
            "README、使用指南、数据字典和实验复现命令与代码同目录维护。",
        ],
        [
            "export.ts 统一列定义和 CSV/JSON 下载。",
            "冻结 manifest 记录输入文件、配置、时间、来源和 SHA-256。",
            "第一阶段核心模型在第二阶段开发后不再随意修改，除非明确回到第一阶段任务。",
        ],
        "形成了可供 INT、机器学习、外部验证和 ns-3 夹具复用的稳定真值接口，并把“仿真产生的数据”和“遥测推断的数据”从数据治理上隔离。",
        [["src/simulation/export.ts", "导出", "nodes/links/tasks/routes/metadata"], ["DATASET_SCHEMA_GUIDE.md", "数据字典", "输出列和物理意义"], ["reports/experiment14b-*/freeze-manifest.json", "冻结清单", "输入哈希与依赖完整性"]],
        boundary="当前工作树仍有较多未提交修改和未跟踪实验产物；正式投稿前需要创建冻结 tag、清洁复现包并记录确切 commit。",
    )

    # Chapter 3
    add_heading(doc, "第 3 章  第二阶段：INT、Ground OAM 与非全知边界", level=1, page_break=True)
    add_improvement(
        doc,
        "3.1 原生 INT 采集链与两种基本模式",
        "第一阶段仪表盘能直接看到全网真值，但真实遥测系统不能全知。第二阶段必须复现 probe 如何经过卫星、每跳写入什么 metadata、sink 如何生成 report，以及 Ground OAM 最终收到什么。",
        [
            "traffic-int 只在业务包经过的路径上采集状态，开销较低但观测受业务分布限制；probe-int 主动规划探测路径，可提高全网覆盖，但产生额外 probe 和 report 开销。",
            "一条 probe 路径由 source、sink 和 hop 序列组成。每跳可以记录节点 CPU、队列、电量/模式，以及相邻链路利用率、容量、时延和 up/warning/down 状态。sink 汇总后沿 reporting path 回传。",
        ],
        [
            "buildIntTelemetryRun 支持浏览器实时演示；离线工具输出 probe plan、hop events、report events 和覆盖率。",
            "probe/header/report/逐跳 metadata 字节分别核算，不能只统计选中路径数量。",
            "同一时间片的直接观测集合 Ωt 是后续所有补全方法的公平输入。",
        ],
        "项目完成了算法级 INT 机制与状态采集语义复现，能够生成逐时间片、逐 probe、逐 hop、逐 report 的可下载数据。",
        [["src/simulation/intTelemetry.ts", "浏览器演示", "实时 INT 运行模型"], ["stage2-int/tools/probe-int-runner.mjs", "离线执行", "probe/hop/report 生成"], ["stage2-int/tools/telemetry-byte-budget.mjs", "开销核算", "header、metadata、report 和链路字节"]],
        formulas=[(r"L_{pkt}=L_{payload}+L_{INT\ header}+\sum_{h=1}^{H}q_hL_{metadata,h}", "3-1", "逐跳 metadata 使数据包随路径长度增长；q_h=0 表示 forward-only。")],
        boundary="第二阶段是协议机制、路径和状态采集逻辑复现，不是 P4/Tofino 硬件流水线，也不是所有业务包逐包仿真；逐包可执行性由实验 13 部分补强。",
    )

    add_improvement(
        doc,
        "3.2 Ground OAM：从报告到全网状态图",
        "INT 报告是离散路径上的局部观测，不等于全网状态。需要一个地面管理实体统一去重、对齐时间片、维护状态年龄、标识 unknown，并调用重构器形成 OAM 视角的全网快照。",
        [
            "OAM 是 Operations, Administration and Maintenance。Ground OAM 接收 sink report，按 satellite_id/link_id 和 slice_index 建立已观测状态表；对未观测对象保留 unknown、confidence 和 AoI。",
            "OAM 重构不是修改第一阶段真值，而是根据已交付观测、历史 OAM 和补全模型估计缺失状态。第一阶段真值只在实验结束后计算 MAE/准确率。",
            "报告未交付意味着 OAM 不可见，即使 probe 在星上已经采集到 metadata，也不能被算作地面已观测。",
        ],
        [
            "按时间片输出 observed nodes、observed links、unknown、confidence、coverage 和 reconstructed snapshot。",
            "对重复报告做一致性/冲突记录，对过旧观测增加 AoI。",
            "检验页把 OAM 结果与真值对照，过程页不显示隐藏真值。",
        ],
        "把“采集到 metadata”和“地面形成全网感知”区分开，使报告丢失、回传中断、未知状态和补全误差都能被独立评价。",
        [["stage2-int/tools/ground-oam-reconstructor.mjs", "OAM 重构", "报告接收、unknown、confidence 和快照"], ["stage2-int/tools/reporting-interruption.mjs", "回传中断", "报告失败和延迟"], ["src/components/IntTelemetry*", "遥测页面", "OAM 视角与后验检验"]],
        formulas=[(r"AoI_o(t)=t-\max\{\tau\leq t:o\in\Omega_\tau\}", "3-2", "对象 o 距离上次直接观测的状态年龄。")],
        boundary="Ground OAM 在仿真中是逻辑地面控制平面，尚未建模真实多地面站控制协议、认证和跨站一致性。",
    )

    add_improvement(
        doc,
        "3.3 INT 过程可视化与可下载证据包",
        "早期第二阶段只显示结果摘要，无法回答 probe 为什么被选、每跳写入什么、报告如何回传、OAM 如何逐片增加覆盖。为答辩和审稿，需要把遥测过程本身变成可审计产物。",
        [
            "离线实验新增 int-process-visualization.json/md，按 slice 组织 probe、hop、report、OAM reconstruction 和 coverage。manifest 显式引用过程包。",
            "遥测页顶部固定时间片与播放控制；中部二维拓扑只显示已观测/已重构状态；右侧展示 probe/hop/report/OAM 细节；底部显示过程时间线、覆盖率热力图、报告和下载。",
            "未观测对象为灰色；当前 probe 路径高亮；已观测链路按状态着色；真值链路不在过程视角直接显示。",
        ],
        [
            "点击 slice 同步所有视图，点击 probe 高亮路径，点击 hop 展开写入 metadata，点击 report 显示 reporting path。",
            "支持下载 probe plan、hop events、report events、OAM reconstruction、coverage 和过程 JSON。",
            "支持实时演示数据与离线实验过程包两种来源。",
        ],
        "第二阶段从“结果表”升级为完整可解释过程，用户可以逐片追踪全网感知是如何形成的，也能检查 unknown、报告失败和重构置信度。",
        [["stage2-int/tools/process-visualization-writer.mjs", "过程包", "统一离线过程模型"], ["stage2-int/tools/process-visualization-inspector.mjs", "验收", "时间片和事件完整性"], ["src/components/IntTelemetry*", "前端", "过程拓扑、时间线和下载"]],
        boundary="过程可视化忠实呈现当前仿真记录，但不等同于真实 P4 数据平面抓包；实验 13 负责补充逐包语义。",
    )

    # Chapter 4
    add_heading(doc, "第 4 章  原生 INT-MC 迁移到动态 LEO 的问题", level=1, page_break=True)
    add_heading(doc, "4.1 原生 INT-MC 的基本技术路线", level=2)
    add_paragraph(doc, "原生 INT-MC 的核心是选择一部分能够覆盖高信息量链路的探测路径，采集路径上的 INT metadata，再利用低秩矩阵补全恢复未观测项。若链路×时间矩阵 X 由少数共同负载模式驱动，则可近似写为 X≈UVᵀ，只需观测 Ω 中的一部分元素。")
    add_formula(doc, r"\min_{U,V}\;\frac{1}{2}\|\mathcal{P}_{\Omega}(X-UV^\top)\|_F^2+\frac{\lambda}{2}(\|U\|_F^2+\|V\|_F^2)", "4-1", "原生低秩补全：观测误差与因子正则化的联合目标。")
    add_paragraph(doc, "路径选择不能直接挑单条链路，因为 INT 的执行单位是路径。原生方法通常根据历史矩阵的杠杆值、覆盖价值和路径成本选择 probe，使一条路径覆盖多个待测链路，再把其余缺失项交给补全器。")
    add_formula(doc, r"\ell_i=\|U_r(i,:)\|_2^2,\qquad Score(P)=\frac{\sum_{e\in P\setminus\Omega}\ell_e}{C(P)}", "4-2", "高杠杆链路和单位路径成本信息收益。")

    add_heading(doc, "4.2 迁移到 LEO 后出现的主要矛盾", level=2)
    add_table(
        doc,
        ["矛盾", "原生隐含假设", "LEO 中的实际问题", "必须增加的约束"],
        [
            ["拓扑动态", "链路行长期存在", "极区、距离和接触窗口使链路产生/消失", "active mask、拓扑版本、局部修复"],
            ["历史复用", "相似历史可直接使用", "少量变化可能位于热点/回传路径", "风险、自适应阈值、质量门禁"],
            ["规模", "采样比例固定", "1584 星下固定路径上限导致采样稀释", "规模自适应硬预算与代表轮换"],
            ["对象类型", "主要补链路矩阵", "CPU/队列/电量同样是关键 INT 对象", "节点—链路异构状态与多指标评分"],
            ["报告可见性", "采到即视为可用", "report 可能因 SGL 窗口/拥塞丢失", "report delivery 与 Ground OAM 门禁"],
            ["闭环合法性", "可离线使用全矩阵", "部署时不知道真实误差", "无 oracle、不确定性与 AoI 反馈"],
        ],
        widths=[1300, 1900, 3100, 3060],
        font_size=8.2,
        first_col_bold=True,
    )
    add_heading(doc, "4.3 为什么待补全矩阵必须有 LEO 约束", level=2)
    add_bullets(doc, [
        "active mask：不存在或物理 down 的链路不是普通缺失值，不能被补成高利用率；其值和状态必须由拓扑掩码固定。",
        "时间窗口：过长窗口会混合多个拓扑分布，低秩关系失效；在线重构只能使用 t 及以前观测。",
        "拓扑类/版本：同一个 link_id 在不同活动图中的邻接和路径作用不同，历史不能无条件复用。",
        "轨道邻居：同轨相邻、相邻轨道面和相近槽位提供结构先验，但只在当前 active mask 中使用。",
        "观测锁定：真实收到的 INT metadata 是硬证据，补全迭代不得覆盖。",
        "不可用链路固定 down：矩阵缺失表示不知道，topology-down 表示不存在/不可用，二者语义必须分开。",
        "节点矩阵并列：CPU、队列、电量和模式与业务、链路和光照耦合，不能把工作重心只放在链路利用率。",
    ])

    # Chapter 5
    add_heading(doc, "第 5 章  增强 LEO-INT-MC 的算法与工程改进", level=1, page_break=True)
    add_paragraph(doc, "项目早期的增强机制确实曾显得“组合式”：多个开关叠加，难以形成统一论文方法。随后主算法被收敛为四个核心量：拓扑版本 ν、非真值不确定性 U、未来风险 R 和硬预算 B。节点/链路物理先验保留在重构层，不再混入路径规划主公式。")

    add_improvement(
        doc,
        "5.1 拓扑预测与拓扑版本化",
        "LEO 的轨道运动可预测，但当前可用路径不代表未来仍可用。只按当前拓扑规划会在极区/切换前后失效；每片全量 fresh replan 又会造成大规模候选生成和评分开销。",
        [
            "接触预测器读取当前和未来时间片的活动链路，计算下一状态切换、接触窗口长度、极区风险、未来断链风险和拓扑漂移压力。",
            "拓扑版本 νt=(ct,ht,ΔEt,τt) 同时记录拓扑类、精确活动边哈希、相对上一版本的增删边，以及距离显著变化的时间。版本不是单纯分数，而是决定 reuse、repair、fresh 的候选空间。",
            "结构缓存与动态评分分离：活动图、候选路径、路径—边覆盖关系和静态成本可复用；利用率、AoI、OAM 目标和报告风险按片更新。",
        ],
        [
            "predict-contact-plan.mjs 生成滚动接触预测与风险字段。",
            "topology-versioned-active-telemetry.mjs 计算版本目标和信息收益摘要。",
            "int-mc-path-selector.mjs 为 reuse/repair/fresh 生成同预算候选并输出触发/拒绝原因。",
        ],
        "把“用了轨道预测”提升为可复现的版本化决策：拓扑相似时复用结构计算，局部变化时只修复受影响路径，变化大或质量风险高时回退 fresh。实验 12 证明了 local repair 的确定性工作量收益。",
        [["stage2-int/tools/predict-contact-plan.mjs", "约 192–614", "工程参数、接触窗口和未来风险"], ["stage2-int/tools/topology-versioned-active-telemetry.mjs", "约 90–259", "版本目标、信息收益和规划成本"], ["stage2-int/tools/int-mc-path-selector.mjs", "约 6186–6600", "结构缓存、三类候选和统一执行"]],
        formulas=[(r"\nu_t=(c_t,h_t,\Delta E_t,\tau_t),\qquad \Delta E_t=E_t\triangle E_{t-1}", "5-1", "拓扑版本及活动边对称差。"), (r"J(E_t,E_k)=\frac{|E_t\cap E_k|}{|E_t\cup E_k|}", "5-2", "拓扑相似度只是一项输入，最终阈值还受风险和 OAM 压力调节。")],
        boundary="当前 48 片正式证据主要证明 local repair，而不是 exact reuse；跨强动态压力和所有星座的统计收益仍未完全站稳。",
    )

    add_improvement(
        doc,
        "5.2 自适应拓扑复用阈值与局部失效",
        "固定 0.94/0.95 的相似度阈值导致 Telesat 长窗口几乎不触发复用；简单降低阈值又可能漏掉关键变化链路。需要让阈值随质量风险、OAM 压力和历史修复成功率变化，并只失效真正受 ΔE 影响的路径。",
        [
            "当 OAM 强制目标多、报告风险高或历史修复失败时，提高复用门槛；当 fresh 规划成本高、变化边比例低且修复历史可靠时，适度降低门槛。",
            "缓存路径 p 只有在 E(p) 与 ΔEt 相交时失效；未受影响路径继续保留。受影响路径删除断边并在局部端点间求替代最短路，而不是清空整个计划库。",
            "复用后执行快速非劣门禁：mandatory target、活动链路覆盖、预计信息收益、报告可交付和硬预算。失败则增量补齐或回退 fresh。",
        ],
        [
            "记录 reuse rejected reason，区分相似度不足、cache miss、失效边、mandatory 缺口、reporting path 不可用等。",
            "将相似度阈值由固定常数改为每片风险/OAM 压力函数。",
            "实验 12 自适应版本在 Telesat 三个 48 片窗口中触发 65/144，全部为局部修复。",
        ],
        "中型星座从 0 触发推进到 45.14% 触发率，并在三个自然窗口均降低规划工作量，同时保持预注册质量门禁。",
        [["stage2-int/tools/int-mc-path-selector.mjs", "结构缓存与 repair 区域", "部分失效、修复和拒绝原因"], ["stage2-int/tools/topology-versioned-risk-planner.mjs", "模式选择区域", "同预算方案比较与质量门禁"], ["reports/experiment12-adaptive-reuse-48slice-generalization", "正式结果", "三窗口 48 片推广性实验"]],
        formulas=[(r"\tau_t=\tau_0+\alpha R_t^{quality}-\beta C_t^{planning},\qquad P_t=P_{valid}^{cache}\cup P_{incremental}", "5-3", "风险提高门槛，规划成本提高复用倾向；计划由有效缓存与增量修复组成。")],
        boundary="当前证据限定于 Telesat 351、三个自然窗口和 0% 额外动态压力；能量平均上升约 0.69%，只能称非劣，不能称能耗下降。",
    )

    add_improvement(
        doc,
        "5.3 无 oracle 不确定性感知闭环",
        "增强版曾在实验脚本中把补全值与第一阶段真值之差直接生成复测目标，这属于 oracle-assisted 泄漏；路径选择器也曾读取隐藏全网 CPU、队列、电量和利用率决定测哪里。部署型 OAM 不知道自己真实误差，必须只用已观测历史和可预测轨道。",
        [
            "不确定性由预测方差、AoI、不同补全模型的分歧、报告冲突和拓扑风险形成，不再使用 |x̂−xtruth|。",
            "第一阶段真值文件从 planner 输入中移除；current utilization、CPU、energy 等信息只能来自上一片 Ground OAM 或已交付 report。轨道位置和未来 contact 仍可使用，因为它们是可预测/公开输入。",
            "反馈至少滞后一片：t 片 OAM 结果只能影响 t+1 及以后 probe；真值只在 probe 与重构全部完成后进入 evaluator。",
        ],
        [
            "删除 high-simulation-validation-error 作为正式 planner 反馈。",
            "实验 7 对 planner、feedback、OAM 和输出字段做无真值泄漏审计。",
            "实验 10–12 的正式增强结果使用 oracle-free replay。",
        ],
        "修复了最严重的实验合法性问题：增强方法不再提前知道哪里错、哪里拥塞或哪里电量低，实验结果可解释为可部署信息边界下的算法表现。",
        [["stage2-int/tools/int-mc-feedback.mjs", "反馈生成", "置信度、冲突、AoI 和模型分歧"], ["stage2-int/tools/int-mc-path-selector.mjs", "输入边界", "OAM/预测输入而非隐藏真值"], ["scripts/runExperiment7*", "实验 7", "oracle-free 审计与门禁"]],
        formulas=[(r"U_o(t)=w_1\sigma_o^2(t)+w_2AoI_o(t)+w_3D_o^{model}(t)+w_4C_o^{conflict}(t)+w_5R_o^{topology}(t)", "5-4", "不确定性完全由可部署信号组成。")],
        boundary="无 oracle 解决了实验公平性，不自动保证估计准确；不确定性质量仍需通过误差—置信度校准和多种子实验验证。",
    )

    add_improvement(
        doc,
        "5.4 风险感知与断链前主动观测",
        "只根据当前不确定性采样会在即将断开的链路上留下长时间状态空洞。LEO 中极区、接触窗口、太阳规避和 SGL 回传中断具有可预测风险，应该在对象消失前提高其观测价值。",
        [
            "风险 R 同时包含未来 down 概率、进入极区概率、reporting failure 概率和拓扑切换概率。风险并不越过预算，而是在预算内改变候选对象优先级。",
            "即将断开的对象若同时具有高 AoI/高不确定性，获得更高信息价值；稳定且刚观测过的对象优先降为 compact 或不写 metadata。",
            "太阳规避、天线占用和星地窗口由第一阶段可预测物理上下文提供，业务拥塞风险只来自历史 OAM。",
        ],
        [
            "接触预测输出 next_transition_time、contact_scarcity、polar_risk 和 report_failure_risk。",
            "统一 planner 在动作价值中加入 risk multiplier，并输出断链前最后观测率。",
            "动态压力实验把参考计划失效和报告中断独立测量。",
        ],
        "系统由“缺失发生后再补救”转向“断链前先测一次”，能够降低切换窗口的长 AoI 和局部盲区。",
        [["stage2-int/tools/predict-contact-plan.mjs", "约 443–614", "未来切换与风险"], ["stage2-int/tools/topology-versioned-risk-planner.mjs", "约 762–893", "风险加权边际价值"], ["scripts/runExperiment8*", "实验 8", "动态性与报告中断敏感性"]],
        formulas=[(r"R_t(o)=r_1P_{down}(o)+r_2P_{polar}(o)+r_3P_{report\ failure}(o)+r_4P_{transition}(o)", "5-5", "未来可用性与回传风险的统一表示。")],
        boundary="风险来自仿真轨道和历史 OAM，不包括真实运营商未公开的临时调度、设备故障和控制命令。",
    )

    add_improvement(
        doc,
        "5.5 单位字节边际信息增益与冗余抑制",
        "只选最长路径会造成 header 膨胀和重复测量；只选最短路径又可能覆盖不到高价值区域。算法应评价一条候选路径在已选集合 S 之后还能新增多少信息，而不是简单累计单链路分数。",
        [
            "对象价值由不确定性 U、风险 R、关键度和 metadata 完整度 q 共同决定；已经被 S 覆盖或与已选对象高度相似的候选受到 redundancy 惩罚。",
            "动作 a=(p,m) 同时包含路径 p 与 metadata 模式 m。full、compact 和 forward-only 对应不同 q 和字节成本。",
            "每轮选择单位字节边际收益最高的动作，更新已观测集合和冗余，再继续选择，直到硬预算无法容纳新动作。",
        ],
        [
            "路径选择器输出 marginal_information_gain、redundancy_penalty、novelty_ratio 和 information_per_kb。",
            "高价值节点在不同轨道面轮换，避免总是测同一组核心节点。",
            "低价值基础路径可被高价值修复路径在固定总预算内替换，不再在基础计划后无上限追加。",
        ],
        "主动测量从“覆盖越多越好”转为“单位成本获得多少新信息”，为覆盖率—开销和误差—开销 Pareto 分析提供了明确算法依据。",
        [["stage2-int/tools/topology-versioned-risk-planner.mjs", "约 762–1202", "硬预算边际贪心"], ["stage2-int/tools/int-mc-path-selector.mjs", "约 3025–3309/6200+", "路径候选、信息量和统一执行"], ["stage2-int/tools/importance-aware-telemetry.mjs", "重要性模块", "节点/链路重要度与轮换"]],
        formulas=[(r"\Delta V_t(a\mid S)=\sum_{o\in Obs(a)\setminus Obs(S)}q_m(o)U_t(o)(1+\alpha R_t(o))-\gamma Redundancy(a,S)", "5-6", "候选动作在当前集合之后的边际价值。"), (r"a^*=\arg\max_{a\notin S}\frac{\Delta V_t(a\mid S)}{C_{bytes}(a)}", "5-7", "单位字节边际收益选择。")],
        boundary="贪心算法是预算背包的工程近似，不保证全局最优；其优势需要在相同候选集和相同预算下与随机、最短路及强规划基线比较。",
    )

    add_improvement(
        doc,
        "5.6 真正的硬字节预算与规模自适应",
        "早期 critical override、OAM 强制复测和极区风险可能越过总字节上限，导致所谓“固定预算比较”被悄悄破坏；另一方面，所有规模使用固定路径上限会在 1584 星下产生严重采样稀释。",
        [
            "硬预算是不可被收益抵消的约束：任何 OAM、极区或高风险对象只能在预算内提高优先级，不能突破 ΣCbytes≤Bt。",
            "规模自适应预算根据目标覆盖、星座规模、路径长度和标准 metadata 成本推导；原生与增强方法使用同一预算口径，以免增强组获得更多隐含采样。",
            "预算不足时执行替换：移除单位价值最低的基础动作，再容纳高价值修复动作；不是在基础计划后追加。",
        ],
        [
            "topology-versioned-risk-planner 在无正硬预算时直接报错，并记录 budget violation。",
            "telemetry-byte-budget 统计 probe base、target mask、hop metadata、report 和实际 ISL 字节。",
            "实验 10–12 将硬预算违例作为零容忍公平门禁。",
        ],
        "确保了“更高质量是否来自更多字节”的问题可以被审计，并缓解大规模固定路径上限造成的失真；同时揭示 Starlink 增强方案在部分实验中仍会使用更多字节，这一结果被如实报告。",
        [["stage2-int/tools/topology-versioned-risk-planner.mjs", "794–1296", "硬预算选择与 violation 记录"], ["stage2-int/tools/scale-adaptive-telemetry-budget.mjs", "约 156–249", "规模预算推导和 admission"], ["stage2-int/tools/telemetry-byte-budget.mjs", "1–80", "实际字节组成核算"]],
        formulas=[(r"\sum_{a\in S_t}C_{bytes}(a)\leq B_t", "5-8", "硬预算是约束，不是可被质量收益抵消的软惩罚。")],
        boundary="规模自适应解决了固定上限稀释，但不能保证所有规模单位节点字节相同；路径长度、覆盖结构和报告数量仍会造成合理差异。",
    )

    add_improvement(
        doc,
        "5.7 选择性 metadata：full、compact 与 forward-only",
        "路径为了串联高价值节点，常常必须经过不重要的中继卫星。若每个中继都写完整 metadata，路径覆盖效率会被无关字段开销抵消。",
        [
            "full 写入全部节点/链路字段，适用于高不确定性或强制目标；compact 保留重构需要的核心字段并减少字节；forward-only 只转发 probe，不暴露该跳节点或本地链路状态。",
            "forward-only 不是把统计字节改小，而是生成可执行的逐跳 target mask；实验 13 在 ns-3 中实际按该 mask 控制 header 增长。",
            "为了避免重要状态被错误略过，只有节点及其本地链路均低价值、低 AoI 且不属于 OAM mandatory 时才允许 forward-only。",
        ],
        [
            "路径选择器为每条路径生成 metadata action，并核算 target-mask 字节和实际字段字节。",
            "报告中分别统计 compact hop、target metadata hop、transit hop 和节省字节。",
            "重要节点按轨道面代表轮换，避免长期 forward-only 造成无限 AoI。",
        ],
        "在不改变路径连通性的情况下减少逐跳写入，特别适合“用短路径串起少数重要对象”的主动测量；小/中规模 pilot 中降低了字节并改善最大 AoI。",
        [["stage2-int/tools/int-mc-path-selector.mjs", "约 5875–5945/6606–6762", "metadata action 与逐跳 mask"], ["stage2-int/tools/importance-aware-telemetry.mjs", "重要性选择", "对象优先级、轮换与 AoI"], ["stage3-system-validation/ns3/scratch/leo-int-system-validation.cc", "逐包验证", "header 增长与 forward-only 执行"]],
        boundary="选择性字段必须和数据平面实现能力匹配；当前 ns-3 证明包级语义，尚未完成 BMv2/P4 程序的真实交换机流水线验证。",
    )

    add_improvement(
        doc,
        "5.8 Ground OAM 强制复测的弱化与保守门控",
        "OAM 强制复测能防止低置信度对象长期失明，但如果每个低置信对象都变成强制目标，会在中小规模造成 500 B/节点/片级开销，甚至使增强方案成为负优化。",
        [
            "复测目标由置信度、AoI、冲突、关键度和拓扑风险共同排序，低价值、短 AoI 的对象不强制；复测只能参与预算内替换。",
            "节点状态耦合和链路张量耦合使用保守门控：只有预测压力、上下文质量和多指标评分同时通过时才修正，避免为了单一利用率改善伤害 CPU/电量/队列。",
            "质量非劣门禁预注册为 CPU、队列、电量、链路利用率 MAE 退化不超过 1% 或 2%（按实验协议），不通过则回退更保守动作。",
        ],
        [
            "删除 OAM critical override 的越界能力，仅保留优先级提升。",
            "对中小规模调低无效复测密度，对大规模保留真正高风险对象。",
            "实验 2 重跑后比较每个指标与字节变化，不只看补全利用率。",
        ],
        "把 OAM 复测从无条件加路径改成预算内、跨指标保守修复，使 Iridium/Telesat 在降低字节的同时显著改善节点和链路误差；Starlink 质量明显改善，但字节仍有上升。",
        [["stage2-int/tools/int-mc-feedback.mjs", "复测候选", "置信度/AoI/冲突"], ["stage2-int/tools/int-mc-path-selector.mjs", "mandatory admission", "预算内替换和门控"], ["scripts/runExperiment2IntMcEnhancementComparison.mjs", "实验 2", "原生/增强统一重跑"]],
        boundary="保守门控降低负优化风险，但并不意味着所有指标必然改善；正式结论应采用 Pareto/条件优势，而不是“全面优于”。",
    )

    add_improvement(
        doc,
        "5.9 统一 reuse、repair 与 fresh 的规划目标",
        "若复用由相似度阈值单独触发、风险由另一个开关修补、预算再由第三个控制器截断，论文方法会像规则拼盘。项目因此把三类规划方式放入同一目标，在相同硬预算下比较。",
        [
            "对 z∈{reuse,repair,fresh} 分别生成候选方案 Sz，使用相同的对象 U、R、metadata action 和硬预算完成选择。",
            "方案价值由预算内信息收益减去规划成本组成；reuse 的规划成本最低，repair 次之，fresh 最高，但低成本不能抵消质量风险。",
            "最终 planner 的公开配置收敛为 byte budget、risk weight、redundancy weight、planning cost weight 和 prediction horizon，历史细开关只用于消融。",
        ],
        [
            "topology-versioned-risk-planner 提供统一 hard-budget selection engine。",
            "路径选择器按 mode prefilter、candidate source 和质量门禁执行。",
            "实验 12 的消融把拓扑版本、风险、局部修复和完整目标拆开。",
        ],
        "主算法从十几个启发式开关收敛为可写进论文的统一决策流程：版本决定候选方式，风险和不确定性决定信息价值，硬预算决定可执行集合。",
        [["stage2-int/tools/topology-versioned-risk-planner.mjs", "约 1206–1296", "三模式候选与硬约束输出"], ["stage2-int/tools/topology-versioned-active-telemetry.mjs", "约 103–259", "统一目标诊断"], ["stage2-int/planning/TOPOLOGY_VERSIONED_ACTIVE_TELEMETRY.md", "方法文档", "正式算法与配置"]],
        formulas=[(r"z_t^*=\arg\max_{z\in\{reuse,repair,fresh\}}\left[\sum_{a\in S_t^z}\Delta V_t(a)-\lambda_p C_{planning}(z,\nu_t)\right]", "5-9", "三种规划方式在相同硬预算下统一比较。")],
        boundary="当前实现仍包含若干工程保护参数；论文应把它们放入实现细节/消融，而不是都宣称为独立创新。",
    )

    # Chapter 6
    add_heading(doc, "第 6 章  重构后端、强基线与多指标探索", level=1, page_break=True)
    add_improvement(
        doc,
        "6.1 LEO 约束下的低秩矩阵补全",
        "原生低秩补全把矩阵中所有空白都视为同类缺失，可能把拓扑 down 链路补成高利用率，也可能改写已观测值。增强重构首先建立结构先验和活动掩码，再执行低秩更新。",
        [
            "初始化按相邻时间片、轨道张量邻居、同类空间分组、链路历史和全局均值逐级回退。初始化已经给每个缺失项一个估计，但它只是局部可解释先验，不是最终全局补全。",
            "低秩步骤在初始化基础上寻找跨链路、跨时间的少数公共模式，降低局部均值带来的偏置。每次迭代后重新锁定直接观测值，并把 inactive 位置固定为 0/down。",
            "初始化与低秩补全分两步的原因是：前者让优化从合理卫星结构出发，后者利用全局相关性；二者分别解决冷启动和全局一致性。",
        ],
        [
            "构建 link×slice 和 node×slice 矩阵、activeMask 与 observedMask。",
            "低秩近似后调用 lockObservedEntries，确保观测不变、inactive 归零。",
            "输出 inferred-only MAE/RMSE/P95，避免直接观测值稀释补全误差。",
        ],
        "使 INT-MC 能正确区分 unknown 与 topology-down，并同时重构节点和链路状态；仍保留原生路径作为增强前基线。",
        [["stage2-int/tools/int-mc-reconstructor.mjs", "约 1267–1400/1880+", "低秩近似、mask 和观测锁定"], ["stage2-int/tools/int-mc-constraints.mjs", "约束模块", "active/time/topology/orbit/down 约束"], ["stage2-int/tools/int-mc-observability.mjs", "可观测性", "直接/推断区域区分"]],
        boundary="低秩假设并非在所有指标和所有压力下成立，因此项目增加强基线并报告无普适赢家。",
    )

    add_improvement(
        doc,
        "6.2 节点—链路异构状态耦合与物理一致性",
        "INT-MC 最初主要补链路利用率，但项目目标是每时间片全网节点与链路状态。CPU、队列、电量、节点模式、链路利用率和状态存在因果/约束关系，完全独立补全会出现高 CPU 却无任务、低能量却高性能模式等违反。",
        [
            "节点状态矩阵 XV=[CPU,Queue,Energy,Mode]，链路状态矩阵 XE=[Utilization,Latency,Loss,Status]。直接观测损失、时间连续性、轨道邻居、节点—链路耦合和物理投影共同约束估计。",
            "CPU 与源任务、过境转发和遥测负载相关；队列与入流、出流和容量相关；电量遵守光照与负载传播；down 链路不能承载正利用率。",
            "为防止错误先验传播，耦合采用保守门控和多指标评分，只有上下文证据强、预计能改善多项误差时才应用。",
        ],
        [
            "metricTensorCoupling、nodeStateCoupling、jointStateCoupling 分开记录应用样本和压力。",
            "energy physics prior 使用 t−1 OAM、电源参数和可预测太阳上下文，不读取 t 真值。",
            "physical consistency 模块统计投影前后的违反数量。",
        ],
        "增强方案不再只优化利用率：实验 2 中三种规模的 CPU、队列、电量和链路利用率 MAE 均较原生下降，节点模式/链路状态准确率保持或改善。",
        [["stage2-int/tools/int-mc-reconstructor.mjs", "约 3173–3455/3800+", "轨道图、节点—链路耦合和保守门控"], ["stage2-int/tools/int-mc-energy-prior.mjs", "电量先验", "因果电池传播"], ["stage2-int/tools/int-mc-physical-consistency.mjs", "物理投影", "范围和状态一致性"]],
        formulas=[(r"\min_{\hat X^V,\hat X^E}L_{obs}+\alpha L_{temporal}+\beta L_{orbit}+\gamma L_{node-link}+\delta L_{physics}", "6-1", "节点—链路异构状态联合约束目标。")],
        boundary="当前多数实现仍是分指标补全后保守耦合，而不是端到端异构图神经网络；论文应准确称为结构/物理约束联合重构。",
    )

    add_improvement(
        doc,
        "6.3 多种可独立调用的补全后端",
        "只与原生低秩方法比较不足以排除“换成熟补全模型也能做到”的质疑。项目保留原方法，并新增 SoftImpute、Kalman 平滑、图邻居、图时序正则、ST-GNN、CoSTCo 和联合 CP 张量后端，统一输入同一 observed mask。",
        [
            "SoftImpute 对奇异值做软阈值，控制有效秩；Kalman/时间平滑对每条 active segment 做状态空间估计；图邻居插值利用同轨和邻轨邻居；图时序正则同时融合低秩、图、时间和先验。",
            "ST-GNN 使用轨道邻居聚合与时间特征学习预测权重；CoSTCo 采用可学习的低秩因子和上下文特征。两者已脱离纯演示轻量参数，但仍是项目内研究实现，不应冒充原论文官方代码。",
            "所有后端锁定直接观测和 inactive mask，并输出推断区误差、运行时间和参数诊断。",
        ],
        [
            "通过 --completion-backend 独立选择后端，不替换原生 low-rank。",
            "正式实验 3 覆盖 Telesat 351 与 Starlink 1584、3 个路径观测种子、4 个节点/链路指标和 48 个时间片；每个星座/seed 内 7 个后端共享 probe plan、节点/链路 observed mask 与实际遥测字节。",
            "以原生低秩为配对参考，按 seed 保留 48 片序列结构，采用 block length=4、4000 次的 moving-block bootstrap，避免把相邻轨道时间片误当作完全独立样本。",
            "统计实验中 Telesat CPU 的 ST-GNN MAE 比低秩下降 10.17%，Telesat 队列的离线 Kalman/RTS 下降 13.26%，Starlink 队列下降 54.70%；电量和多数 Starlink 指标仍由低秩取得最低 MAE。",
        ],
        f"补强了论文的算法对照：正式产物的公平性门禁{'全部通过' if exp3_fairness_passed else '未全部通过'}，共有 {exp3_significant_count} 个“星座×指标×后端”组合的配对 95% CI 完全低于 0。结果同时证明项目不依赖单一补全器，并否定了“某个成熟后端在所有指标普遍最强”的过度主张。",
        [["stage2-int/tools/int-mc-additional-completion-backends.mjs", "约 97–480", "SoftImpute、Kalman、图邻居和图正则"], ["stage2-int/tools/int-mc-reconstructor.mjs", "约 1320–1880", "后端路由、ST-GNN 和 CoSTCo"], ["scripts/runExperiment3StatisticalStrongBaselines.mjs", "正式统计 runner", "同 mask、多种子、多指标配对实验"], ["scripts/experiments/experiment3StrongBaselineStats.mjs", "统计工具", "哈希公平审计与移动块 bootstrap"], ["reports/experiment3-statistical-strong-baselines", "实验 3 正式产物", "HTML、CSV、JSON 与逐片证据"]],
        formulas=[(r"\hat X=\arg\min_Z\frac{1}{2}\|\mathcal{P}_\Omega(X-Z)\|_F^2+\lambda\|Z\|_*", "6-2", "SoftImpute 的核范数正则形式。"), (r"\hat X=\arg\min_Z L_{obs}+\lambda_g\mathrm{tr}(Z^\top L_GZ)+\lambda_t\|ZD_t\|_F^2+\lambda_*\|Z\|_*", "6-3", "图—时间—低秩联合正则的概念目标。")],
        boundary="三个 seed 改变的是路径级 INT 观测 mask，物理真值仍来自同一个 48 片轨道/业务窗口，因此证明的是观测选择稳健性而非跨历元泛化。Kalman/RTS 使用窗口内未来已交付观测，只能作为离线非因果参考；ST-GNN 与 CoSTCo 是项目内研究实现，不是原论文官方代码。",
    )

    add_improvement(
        doc,
        "6.4 多指标三维张量联合补全探索",
        "逐指标二维补全忽略 CPU、队列、电量等指标之间的隐式关系。项目探索把时间×实体×指标组织为三维张量，以 CP 分解学习共享潜因子，并在相同 INT 观测和字节下比较。",
        [
            "节点张量 XV∈R^{T×N×Kv}，链路张量 XE∈R^{T×E×Ke}。plane_id 和 slot_id 不是通过随意加乘作为数值特征，而是作为离散结构、邻接或嵌入，避免产生不存在的序数距离。",
            "不同量纲先只用训练/直接观测统计归一化；观测 mask 按指标保存，直接观测锁定；拓扑 down 链路由 hard mask 排除。",
            "CP 预测可再做物理投影，但必须分开报告“张量本身收益”和“投影收益”，否则中大型改善可能只是投影贡献。",
        ],
        [
            "实现 joint-cp 与 joint-cp-physics 两个后端，记录 rank、epochs、参数量、observed loss 和 wall time。",
            "实验保持相同 observed mask 与 telemetry bytes，比较二维逐指标和多指标联合补全。",
            "由于现有证据未证明 CP 在各规模稳定降低所有指标，已将其降为辅助探索，不偏离主动遥测主线。",
        ],
        "项目具备多指标联合研究能力，也明确识别了物理投影混淆因素；这项工作可作为后续扩展，而非当前论文最核心主张。",
        [["stage2-int/tools/int-mc-joint-tensor-completion.mjs", "约 77–结尾", "联合 CP 张量与观测训练"], ["stage2-int/tools/int-mc-reconstructor.mjs", "约 2530–2550/3970+", "后端接入和物理投影"], ["reports/experiment3-joint-tensor-completion", "探索实验", "二维与联合张量比较"]],
        formulas=[(r"\mathcal{X}_{t,i,k}\approx\sum_{r=1}^{R}A_{t,r}B_{i,r}C_{k,r}", "6-4", "时间、实体和指标三维 CP 分解。")],
        boundary="现阶段尚不能声称多指标 CP 全面优于二维补全；中大型收益可能主要来自物理投影，需更严格消融后才能成为独立创新。",
    )

    add_figure(doc, create_experiment3_baseline_figure(), "图 4  实验 3：中大型星座多种子、多指标强补全基线", width=Inches(6.15))

    # Chapter 7 metrics
    add_heading(doc, "第 7 章  实验方法、评价指标和公平性原则", level=1, page_break=True)
    add_heading(doc, "7.1 公平实验的统一协议", level=2)
    add_numbered(doc, [
        "冻结第一阶段输入、配置、轨道快照、业务数据和随机种子；各方法共享同一真值。",
        "同一对比中的方法使用相同候选拓扑、相同时间片、相同 observed mask 或相同硬字节预算。",
        "planner 仅使用过去 OAM、已交付报告和可预测轨道；真值只进入后验 evaluator。",
        "直接观测值不参与 inferred-only 误差的稀释；topology-down 从缺失集合中排除。",
        "同时报告质量与成本，保留退化和不显著结果，不按单一最好指标选择结论。",
        "大型实验采用逐场景独立进程、流式时间片和固定并发，避免 16 GB 设备换页导致测量偏差。",
    ])
    add_heading(doc, "7.2 重构质量指标", level=2)
    add_formula(doc, r"MAE=\frac{1}{|\mathcal{M}|}\sum_{(i,t)\in\mathcal{M}}|\hat x_{i,t}-x_{i,t}|", "7-1", "仅在待推断集合 M 上计算的平均绝对误差。")
    add_formula(doc, r"RMSE=\sqrt{\frac{1}{|\mathcal{M}|}\sum_{(i,t)\in\mathcal{M}}(\hat x_{i,t}-x_{i,t})^2}", "7-2", "RMSE 对大误差更敏感。")
    add_formula(doc, r"P95AE=Q_{0.95}(|\hat x-x|),\qquad Accuracy=\frac{N_{correct}}{N_{evaluated}}", "7-3", "尾部绝对误差与离散状态准确率。")
    add_table(
        doc,
        ["指标", "物理含义", "为何不能只看它"],
        [
            ["CPU MAE", "重构 CPU 利用率与真值的平均百分点评估误差", "低均值可能掩盖热点节点尾部"],
            ["队列 MAE", "任务/转发队列深度误差", "需要同时看 P95 与拥塞事件召回"],
            ["电量 MAE", "SoC 百分点误差", "大规模中慢变量历史偏差可能积累"],
            ["节点模式准确率", "normal/power-saving/offline 分类正确率", "类别不平衡时应结合每类召回"],
            ["链路状态准确率", "up/warning/down 分类正确率", "inactive mask 不能被算作普通正确样本"],
            ["链路利用率 MAE", "全部评价链路的利用率误差", "直接观测可能稀释真正补全误差"],
            ["补全利用率 MAE", "只在未直接观测活动链路上的误差", "仍需 RMSE/P95/高利用率 F1"],
        ],
        widths=[1800, 3850, 3710],
        font_size=8.5,
        first_col_bold=True,
    )
    add_heading(doc, "7.3 遥测与规划开销指标", level=2)
    add_formula(doc, r"O_{B/node/slice}=\frac{B_{probe}+B_{metadata}+B_{report}+B_{mask}}{|V|T}", "7-4", "每节点每时间片遥测字节是归一化通信成本。")
    add_formula(doc, r"C_{plan}=w_1N_{candidate}+w_2N_{shortest}+w_3N_{score}+w_4N_{marginal}+w_5N_{repair}", "7-5", "确定性规划工作量分解，比单次墙钟更稳定。")
    add_bullets(doc, [
        "通信开销：总 INT 字节、每节点每片字节、逐跳 metadata、report、target mask、ISL 实际承载字节。",
        "能量开销：probe/report 的发送、接收和处理能耗；当前用于同模型相对比较，不冒充真实星载硬件功耗。",
        "规划开销：候选路径数、最短路调用、评分重算、边际收益计算、图重建、局部修复和规则下发数。",
        "可用性：report delivery、最大/平均 AoI、mandatory target coverage、probe failure、MTU drop。",
        "单位收益：有效覆盖/KB、误差下降/KB、信息增益/KB，以及 Pareto 前沿而不是单指标排序。",
    ])
    add_heading(doc, "7.4 非劣与因果门禁", level=2)
    add_table(
        doc,
        ["门禁", "典型判据", "意义"],
        [
            ["硬预算", "budget violation = 0", "高风险目标不能越过预算"],
            ["质量非劣", "MAEreuse ≤ 1.02×MAEfresh", "规划收益不能靠漏测换取"],
            ["覆盖非劣", "Coverage reuse ≥ Coverage fresh−1%", "保护活动对象覆盖"],
            ["强制目标", "MandatoryCoverage = 100%", "OAM 关键对象必须被预算内覆盖"],
            ["报告交付", "Deliveryreuse ≥ Deliveryfresh−1%", "不能忽略星地回传"],
            ["时间因果", "反馈最少滞后一片，truth read=0", "排除未来/真值泄漏"],
        ],
        widths=[1600, 3000, 4760],
        font_size=8.5,
        first_col_bold=True,
    )

    # Chapter 8 experiments
    add_heading(doc, "第 8 章  实验 1–14 的完整记录与结论", level=1, page_break=True)
    add_paragraph(doc, "本章按实验编号记录目的、设计、指标、结果、结论与边界。开发期旧版本已归档；正文优先引用当前正式或最新代表性结果。")

    add_experiment(
        doc,
        "1",
        "卫星网络仿真数据真实性校准",
        "证明模型不是“能跑通”就自称真实，而是在公开轨道、物理规律、拓扑约束和外部网络统计下具有可解释一致性。",
        [
            "主规模从小型开发拓扑改为 Starlink 72×22=1584，并冻结轨道/业务/配置。",
            "内部检查轨道周期、速度、高度、plane-slot、链路度数、极区、遮挡、容量、能量守恒、空业务 CPU 和任务时延口径。",
            "外部对照使用 CelesTrak GP/OMM、Cloudflare Radar、RIPE Atlas；对不可获得的 CPU/电量/队列只做方程、范围和敏感性验证。",
            "修复分钟级伪任务时延后，用户侧 RTT 与内部任务完成时间分开评价。",
        ],
        ["23/23 内部物理一致性检查", "轨道/拓扑外部偏差", "Radar 时序相关性", "RIPE RTT P50/P95 比例", "任务状态和时延分布"],
        [
            "电池容量 1200 Wh、峰值太阳翼发电约 762.16 W 等参数链能够被公式复算。",
            "用户侧延迟口径修正后 P50 回到可接受量级；内部任务 P95 仍受拥塞和计算等待影响，不能直接等同于用户 RTT。",
            "Radar 相关系数 0.9134 来自校准/拟合证据，不被当作独立泛化。",
            "真实运营商逐星 CPU、队列、电量不可获得，因此实验结论限定为高可信仿真潜变量。",
        ],
        "实验 1 能证明模型满足公开轨道与研究级物理/业务约束，适合作为 INT 算法同模型真值环境；它不能证明复刻运营商内部状态。",
        "reports/experiment1-satellite-data-authenticity/experiment1-research-report.html",
        "不能以内部约束通过替代外部真实逐点对比，也不能把校准段 Radar 相关性写成盲测泛化性能。",
        formula=(r"S_{overall}=\sum_k w_kS_k,\qquad \sum_kw_k=1", "8-1", "综合评分是分项加权摘要，必须同时展示原始分项，不能把一个分数当作真实性本身。"),
    )

    add_experiment(
        doc,
        "2",
        "原生 INT 基线与增强 LEO-INT-MC 综合对比",
        "建立 traffic-int、full probe-int、shortest-path、random sampling、原生 INT-MC 和增强 LEO-INT-MC 的统一基线，并评价覆盖、开销与全指标重构。",
        [
            "三种规模均运行 48 时间片；增强前后共享第一阶段真值、候选路径和无 oracle 输入。",
            "traffic-int 代表低开销部分观测；full probe-int 代表高覆盖高开销；随机和最短路是弱规划基线。",
            "原生/增强 INT-MC 比较 CPU、队列、电量、模式、链路状态、链路利用率和 inferred-only 利用率，并记录字节和能量。",
        ],
        ["直接/有效覆盖率", "B/节点/时间片", "CPU/队列/电量 MAE", "节点/链路状态准确率", "链路利用率与 inferred-only MAE/RMSE/P95", "AoI 和报告交付"],
        [
            "Iridium：字节 506.2424→437.8505，CPU 0.3533→0.2467，队列 3.8306→2.4470，电量 0.8709→0.1343，链路利用率 2.2232→1.1725。",
            "Telesat：字节 426.4387→399.3001，CPU 2.1281→0.9093，队列 0.0874→0.0243，电量 1.3873→0.3150，链路利用率 0.4525→0.3467。",
            "Starlink：字节 18.6995→23.7407（上升），CPU 1.8723→0.9599，队列 0.1681→0.0844，电量 6.3647→2.1136，链路利用率 2.4329→1.9147。",
            "增强方案不是所有成本都下降：大型星座以约 26.96% 的归一化字节增量换取多个误差显著下降。",
        ],
        "实验 2 支持“增强方案在三规模改善多项 OAM 重构质量，并在小/中规模降低通信字节”的条件结论；不支持“所有规模、所有指标和所有开销全面优于原生”。",
        "reports/experiment2-baseline-comparison-oracle-free-replay/experiment2-comprehensive-baseline-report.html",
        "原生 INT-MC 不是动态 LEO 的唯一直接竞品，因此实验 3 又增加成熟补全基线；Starlink 字节上升必须在正文中披露。",
    )

    if exp2_rows:
        exp2_table_rows = []
        for name in ("Iridium 66", "Telesat 351", "Starlink 1584"):
            before = next(r for r in exp2_rows if r["constellation_short_label"] == name and r["version"] == "before")
            after = next(r for r in exp2_rows if r["constellation_short_label"] == name and r["version"] == "after")
            exp2_table_rows.append([
                name,
                f"{fmt(before['telemetry_bytes_per_node_slice'])}→{fmt(after['telemetry_bytes_per_node_slice'])}",
                f"{fmt(before['cpu_mae'])}→{fmt(after['cpu_mae'])}",
                f"{fmt(before['queue_depth_mae'])}→{fmt(after['queue_depth_mae'])}",
                f"{fmt(before['energy_percent_mae'])}→{fmt(after['energy_percent_mae'])}",
                f"{fmt(before['link_utilization_mae'])}→{fmt(after['link_utilization_mae'])}",
                f"{fmt(before['utilization_inferred_mae'])}→{fmt(after['utilization_inferred_mae'])}",
            ])
        add_table(doc, ["星座", "B/节点/片", "CPU MAE", "队列 MAE", "电量 MAE", "链路利用率 MAE", "补全利用率 MAE"], exp2_table_rows, widths=[1320, 1340, 1260, 1260, 1260, 1470, 1450], font_size=7.8, first_col_bold=True)
        delta_figure = create_experiment2_delta_figure(exp2_rows)
        if delta_figure:
            add_figure(doc, delta_figure, "图 5  实验 2：增强方案相对原生 INT-MC 的全指标误差变化", width=Inches(6.35))

    add_experiment(
        doc,
        "3",
        "中大型星座多种子、多指标强补全基线",
        "在完全相同 INT 观测、observed mask 和遥测字节下，区分低秩、时间、图结构和机器学习补全器本身的效果，并检验结论对路径观测选择的稳健性。",
        [
            "主实验选择 Telesat 351 与 Starlink 1584，每个星座使用 seed 11/23/37 生成三组路径级 INT 观测 mask，每组覆盖 48 个时间片。",
            "比较 low-rank、SoftImpute、Kalman/RTS、graph-neighbor、graph-regularized、ST-GNN 与 CoSTCo；同一 profile/seed 内仅切换 completion backend。",
            "同时评估 CPU 利用率、节点队列深度、电量 SoC 和链路利用率，报告 inferred-only MAE/RMSE/P95、R2、运行时间和直接观测率。",
            "以低秩为配对参考，对每个 seed 的逐片 MAE 序列执行长度 4、4000 次 moving-block bootstrap，并记录胜/平/负。",
        ],
        ["inferred-only MAE/RMSE/P95/R2", "配对 95% CI 与胜/平/负", "路径/mask/字节哈希公平性", "直接观测率", "运行时间与参数规模"],
        [
            f"168 条方法—指标聚合记录和 8064 条逐时间片记录完整生成；6 个 profile/seed 公平门禁{'全部通过' if exp3_fairness_passed else '未全部通过'}。",
            "Telesat CPU：ST-GNN mean MAE=2.7884，相对低秩下降 10.17%；Kalman/RTS、graph-regularized 与 CoSTCo 也达到统计显著改善。",
            "Telesat 队列：离线 Kalman/RTS MAE=0.0715，相对低秩下降 13.26%；Starlink 队列 MAE=0.0177，下降 54.70%，两者 CI 均完全低于 0。",
            "Telesat/Starlink 的电量，以及 Starlink 的 CPU 和链路利用率仍由原生低秩取得最低 MAE；成熟后端并未自动全面胜出。",
            "计算代价差异明显：Starlink 经典后端平均单指标低于 0.5 s，而 ST-GNN/CoSTCo 约为 75.2 s/32.7 s。",
        ],
        "强基线缺口已从“单次、单指标 pilot”补强为中大型、三观测种子、四指标的同预算统计对照。证据支持后端可替换和指标依赖性，也进一步确认论文主创新应放在 LEO 主动遥测规划，而不是宣称某个补全器普遍最强。",
        "reports/experiment3-statistical-strong-baselines/index.html",
        "三个 seed 共用同一第一阶段物理真值窗口，尚不能代替独立轨道历元和独立业务窗口；Kalman/RTS 是离线非因果上界参考，不能作为严格在线部署结果。",
    )

    add_experiment(
        doc,
        "4",
        "增强机制消融实验",
        "逐项移除轨道先验、OAM 反馈、节点耦合、电量先验、张量上下文等机制，判断哪些改动有实质贡献，哪些只增加开销。",
        [
            "完整增强方案作为参照，每个变体只关闭一类机制。",
            "三规模或代表规模共享硬预算、业务、拓扑和 seed。",
            "同时检查误差、字节、能量和质量门禁，保留零效应和负效应。",
        ],
        ["各指标相对完整方案变化", "通信/能量开销", "机制实际应用次数", "门禁通过率"],
        [
            "部分机制在某些规模上没有触发，因此移除后指标几乎不变，却可能减少计算/开销。",
            "节点状态耦合和电量物理先验对大规模电量/CPU 误差更有贡献；过强 OAM 复测会伤害中小规模开销。",
            "单次消融中的小差异不能独立证明因果，实验 11 进一步使用多 seed 和动态压力。",
        ],
        "消融推动项目删除或弱化无效开关，并把正式方法收敛到拓扑版本、风险、信息收益、硬预算四个核心量。",
        "reports/experiment4-* / reports/experiment11-dynamic-equal-budget-ablation",
        "某机制在当前窗口不触发不等于理论无用；必须同时报告 trigger count 和 conditional effect。",
    )

    add_experiment(
        doc,
        "5",
        "遥测开销分解",
        "避免只用“总字节”片面表述开销，分解通信、能量、路径规划、报告与 AoI，识别拓扑复用真正节省了什么。",
        [
            "从一次完整运行中分离 probe base、target mask、hop metadata、report、ISL 承载字节和星地回传。",
            "记录候选生成、最短路、评分、边际收益、局部修复和图重建次数。",
            "比较原生与增强在三个规模上的总成本构成。",
        ],
        ["B/节点/片", "pipeline bytes", "telemetry energy", "selected paths", "planning work", "report delivery/AoI"],
        [
            "Starlink 增强并非更省通信：B/节点/片约 18.70→23.74；早期一版 pipeline/能量也有明显上升。",
            "拓扑复用若仍发送全部旧 probe，只能降低规划计算开销，不能自动降低网络字节。",
            "选择性 metadata、预算内替换和稳定对象历史估计才是通信开销下降的来源。",
        ],
        "实验 5 明确了两类开销：规划复用收益与真实网络通信收益必须分别报告，不能用一个总字节或墙钟替代。",
        "reports/experiment5-*",
        "能量是同模型相对估计，不是星载硬件实测；墙钟受系统噪声影响，确定性计数应作为主证据。",
    )

    add_experiment(
        doc,
        "6",
        "采样率与 Pareto 前沿",
        "寻找误差和遥测开销的合理平衡点，而不是预先固定一个凭感觉的采样率。",
        [
            "对 5%–40% 多个预算档运行原生和增强方法。",
            "横轴为通信开销，纵轴为 CPU/队列/电量/链路利用率误差或综合非劣分数。",
            "识别不被其他点同时在成本与误差上支配的 Pareto 前沿和 knee。",
        ],
        ["各指标 MAE vs B/节点/片", "有效覆盖/KB", "综合质量门禁", "Pareto knee"],
        [
            "Iridium 的原生/增强 knee 约为 20%/40%；Telesat 约 15%/15%；Starlink 约 10%/10%。",
            "不同规模的最优预算不同，证明固定路径上限或固定采样率会导致过采样或稀释。",
        ],
        "实验 6 支持规模自适应预算和多目标控制；最终方案应报告 Pareto/条件优势，而不是强求所有指标全面最好。",
        "reports/experiment6-*",
        "knee 依赖当前业务、时间窗口和重构器，需用独立 seed 检查稳定性。",
        formula=(r"x\prec y\Longleftrightarrow C_x\leq C_y,\ \mathcal{E}_x\leq\mathcal{E}_y,\ (C_x<C_y\ \mathrm{or}\ \mathcal{E}_x<\mathcal{E}_y)", "8-2", "Pareto 支配：一个方案只有在成本和误差均不差且至少一项更好时才支配另一个。"),
    )

    add_experiment(
        doc,
        "7",
        "无真值泄漏与严格遥测边界审计",
        "证明 planner 和 OAM 在决定测哪里与补什么时没有读取第一阶段隐藏真值。",
        [
            "扫描 planner 输入、feedback 记录、OAM 构建和 manifest，禁止 truth/error/current hidden state 字段进入决策。",
            "验证反馈至少滞后一片，未来时间片观测不可见。",
            "真值只允许在 evaluation 输出阶段读取。",
        ],
        ["truth read count", "future observation count", "feedback lag", "observed lock", "manifest 输入白名单"],
        ["正式审计报告 24/24 检查通过；快速回归 fixture 的关键门禁也全部通过。", "此前 high-simulation-validation-error 和路径选择器读取隐藏状态的问题已从正式流程移除。"],
        "实验 7 是后续所有增强结果的合法性基础：当前正式结果属于 oracle-free 仿真，而不是知道答案后的上界。",
        "reports/experiment7-* / npm run test:experiment7",
        "代码审计证明输入边界，不证明 OAM 置信度已完美校准。",
    )

    add_experiment(
        doc,
        "8",
        "动态拓扑、参考计划失效与报告中断敏感性",
        "检验方法在额外断链、路径过期和星地回传中断下是否仍稳定，并量化动态压力对观测和重构的影响。",
        [
            "对 0%、25% 等动态压力注入额外链路失败，多个 seed 运行。",
            "分别统计当前 probe 失败、旧 reference plan 失效和 report delivery 中断。",
            "报告均值、置信区间和质量变化。",
        ],
        ["additional probe failure", "reference-plan failure", "report delivery", "重构 MAE/AoI", "置信区间"],
        [
            "25% 压力下额外 probe failure：Iridium 约 2.97% [2.03,3.92]，Telesat 3.39% [2.95,3.84]，Starlink 15.30% [14.66,15.93]。",
            "同压力下旧 reference plan 失效率约 34.69%、29.78%、19.44%，说明静态计划在动态 LEO 中会快速失效。",
            "20% 回传中断下 report delivery 约 80%，Ground OAM 可用观测和 AoI 随之恶化。",
        ],
        "实验 8 直接说明为什么需要拓扑预测、局部修复和 reporting path 约束，而不能把地面静态 INT-MC 原样搬入 LEO。",
        "reports/experiment8-*",
        "注入故障是受控压力，不等同于真实运营故障分布；结论是敏感性而非实际故障率估计。",
    )

    add_experiment(
        doc,
        "9",
        "外部轨道、Radar 与 RIPE Atlas 初步验证",
        "将第一阶段输出与公开外部来源比较，并明确哪些证据是校准、哪些是验证。",
        [
            "两个独立轨道历元共 3168 颗传播，检查 SGP4 成功率和跨历元位置误差。",
            "Radar 评估宏观业务时序，RIPE Atlas Starlink 探针评估用户侧 RTT。",
            "冻结配置并报告时间差、样本数、语义口径。",
        ],
        ["SGP4 failure", "ECI position MAE", "Radar Pearson", "RTT P50/P95 ratio"],
        [
            "3168 条传播无 SGP4 failure；在平均约 132.98 h 的快照年龄下 ECI MAE 约 174.087 km。",
            "Radar 0.9134 属于校准结果，不是盲测；RIPE 对照中模型 P50 约 23.649 ms、外部约 28.855 ms。",
        ],
        "实验 9 提供了第一轮外部锚点，但暴露出轨道快照年龄和校准/验证混用问题，推动实验 14/14B 的冻结盲测协议。",
        "reports/experiment9-*",
        "外部 RTT 只验证用户侧网络趋势，不能验证内部星间队列、CPU 和能量。",
    )

    add_experiment(
        doc,
        "10",
        "三规模、同预算、多种子动态比较",
        "从单次结果扩展到多 seed，判断增强优势是否稳定，并对所有方法执行公平门禁。",
        [
            "三规模×多个压力/方法×10 seeds，共 270 method rows。",
            "原生与增强共享硬预算、truth、业务和候选环境。",
            "按指标和置信区间把结果分为 positive、negative、uncertain。",
        ],
        ["90/90 fairness", "多 seed 均值/区间", "positive/negative/uncertain 计数", "跨规模稳定性"],
        [
            "90/90 公平检查通过。",
            "Iridium：3 positive、12 negative、6 uncertain；Telesat：15/0/6；Starlink：11/1/9。",
            "增强优势主要在中大型成立，小型存在明显适用边界。",
        ],
        "实验 10 支持把论文核心结论限定为中大型动态 LEO 的条件优势，而不是追求所有规模全面胜出。",
        "reports/experiment10-equal-budget-dynamic-multiseed",
        "positive/negative 数量依赖预注册阈值和指标集合，正文必须同时给出原始效应量。",
    )

    add_experiment(
        doc,
        "11",
        "动态等预算多种子消融",
        "在三规模、0/25% 压力、多 seed 下评估完整增强和各机制变体，检验贡献与交互。",
        [
            "三规模×两种压力×10 seeds×6 variants，60 个配对场景。",
            "关闭轨道预测、节点耦合、电量先验、张量业务上下文等机制。",
            "保持预算和 causal/OAM 输入相同。",
        ],
        ["60/60 fairness", "各机制条件效应", "触发率", "误差/开销非劣", "positive/negative/uncertain"],
        [
            "当前汇总约 29 positive、37 negative、234 uncertain，大量不确定结果说明许多机制在特定场景不触发或效应小。",
            "结果不支持把每一个工程开关都包装成创新；支持收敛核心算法并保留适用边界。",
        ],
        "实验 11 的价值不是得到全正结果，而是识别组合式增强的冗余和风险，为统一算法目标提供依据。",
        "reports/experiment11-dynamic-equal-budget-ablation",
        "大量 uncertain 不能解释为“机制有效”，也不能简单解释为“机制无用”；需要 trigger-conditional 分析。",
    )

    add_experiment(
        doc,
        "12",
        "拓扑版本化、局部修复与复用贡献",
        "独立证明拓扑版本和局部修复是否降低路径规划工作量，并检查这种收益是否靠质量退化或预算越界获得。",
        [
            "先运行 18 个自然场景评估触发率，再用高/中/边界相似度窗口做定向配对。",
            "最新推广实验在 Telesat 351 的 window-00/08/16 各运行完整 48 片，比较 fresh-only、固定阈值与自适应 repair。",
            "主指标采用候选/评分/边际计算等确定性工作量；墙钟仅辅助。",
        ],
        ["trigger rate", "exact reuse/local repair/fresh fallback", "planning work reduction", "score/marginal recomputation", "bytes/energy", "质量非劣"],
        [
            "三个窗口触发 18/48、20/48、27/48，总计 65/144=45.14%，全部为 local repair、exact reuse=0。",
            "平均规划工作量下降 23.037%，评分重算下降 34.663%，边际收益评估下降 19.829%。",
            "平均遥测字节下降 0.454%，能量反而上升 0.690%，三个窗口质量门禁均通过。",
            "早期 Starlink 定向高相似场景可见约 51.37% 条件工作量信号，但自然触发和质量门禁不足，未被写成普适结论。",
        ],
        "可以证明：在 Telesat 351、自然轨道动态、0% 额外压力和三个 48 片窗口中，自适应 local repair 在质量非劣下显著减少确定性规划工作量。不能证明 exact reuse 或所有压力/星座普适。",
        "reports/experiment12-adaptive-reuse-48slice-generalization/ADAPTIVE_REUSE_48SLICE_REPORT.md",
        "通信字节仅小幅下降、能量小幅上升；墙钟不稳定，因此不能写成“规划延迟和所有开销都显著下降”。",
        formula=(r"\mathbb{E}[\Delta]=P(trigger)\,\mathbb{E}[\Delta\mid trigger]", "8-3", "整体收益由触发率和触发后的条件收益共同决定。"),
    )
    add_figure(doc, ROOT / "reports" / "experiment12-adaptive-reuse-48slice-generalization" / "overview.png", "图 6  实验 12：Telesat 351 三个 48 片窗口的自适应局部修复结果", width=Inches(6.3))

    add_experiment(
        doc,
        "13",
        "ns-3 小规模逐包系统级交叉验证",
        "回答聚合模型估计的 INT header、MTU、队列竞争、报告交付和 AoI 在真实数据包执行语义下是否仍成立。",
        [
            "从第一阶段导出 66 星、20 时间片动态窗口，导入链路切换、容量、时延和业务流。",
            "比较 full metadata 与 selective metadata，并加入 MTU/过载压力。",
            "逐包测量实际字节、业务吞吐、队列 P95、丢包、report delivery 和 OAM AoI，再和聚合模型同场景比较。",
        ],
        ["planned/actual telemetry bytes", "MTU drop", "queue P95", "delivery ratio", "AoI", "aggregate-vs-ns3 relative error/trend"],
        [
            "selective 相对 full：计划遥测字节下降 40.86%，report delivery 提升 22.63 个百分点，OAM AoI 下降约 12%。",
            "full 的 MTU drop 均值约 827.67，selective 约 31，说明逐跳字段控制具有包级意义。",
            "核心场景聚合模型与 ns-3 的 P95 绝对相对误差约 4.17%，趋势一致率 83.33%；2×过载下 P95 误差约 101.81%，聚合队列模型在极端场景明显失真。",
        ],
        "实验 13 排除了“选择性 metadata 只存在于记账公式中”的主要质疑，并限定了聚合模型在正常压力下的适用范围；极端过载需依赖逐包结果。",
        "reports/experiment13-system-validation/EXPERIMENT_13_SYSTEM_VALIDATION.md",
        "这是 66 星小规模 ns-3 交叉验证，不是 1584 星逐包复现，也不是 P4 硬件实现。",
        formula=(r"U_e(t)=\frac{B_{business}(t)+B_{telemetry}(t)}{C_e\Delta t}", "8-4", "逐包环境中遥测字节真实占用链路容量并与业务竞争。"),
    )
    add_figure(doc, ROOT / "reports" / "experiment13-system-validation" / "experiment13-preview.png", "图 7  实验 13：ns-3 逐包交叉验证概览", width=Inches(6.3))

    add_experiment(
        doc,
        "14",
        "多源外部冻结验证协议",
        "把外部真实性从“选取公开曲线做对照”提升为冻结代码/参数、校准段与测试段分离、测试结果不回调的协议。",
        [
            "冻结当前代码、配置和参数哈希；轨道输入使用 GP0，未来 GP1/SupGP 只做验证。",
            "Radar 前半段校准、后半段测试；RIPE Atlas 与 M-Lab 作为独立网络性能外部源。",
            "报告轨道误差、RTT/吞吐 CDF、地域箱线图、业务时序和置信区间覆盖。",
        ],
        ["轨道 ECI/沿轨/横轨误差", "Radar holdout Pearson/DTW", "RIPE RTT CDF", "M-Lab RTT/throughput ratio", "CI coverage"],
        [
            "第一版实验 14 暴露 GP 快照较旧和验证暴露风险：未来快照 ECI MAE 约 578/615 km；Radar holdout Pearson 约 0.955 但不是研究者盲测。",
            "RIPE RTT 比例约 0.793；M-Lab 指标语义并不完全等价，不能强行合并为同一真值。",
        ],
        "实验 14 最重要的结果是建立外部证据边界和冻结协议，并识别旧快照/非盲测问题，随后由 14B 修正。",
        "reports/experiment14-multisource-external-blind-validation",
        "多个公开源只能验证其公开语义：轨道、用户 RTT、宏观流量和吞吐；不能验证运营商内部 ISL、CPU、电量和队列。",
    )

    add_experiment(
        doc,
        "14B",
        "前瞻性新鲜轨道快照冻结与有界外推",
        "在不持续下载和无限膨胀磁盘的前提下，形成一次内容确实更新、时间因果合法、可复现的新鲜轨道实验窗口。",
        [
            "T0 获取最新 GP0/OMM，检查每星 epoch 年龄，筛选 1584 颗一致轨道群组，冻结来源、时间、配置和哈希。",
            "只在 [T0,T0+H] 内传播 48 时间片/约 4 h；实验结束后不继续更新。",
            "未来时刻获取 GP1/SupGP，仅比较 GP0 外推误差，不反馈修改模型。Radar/RIPE 同样等待真正未来测试段。",
        ],
        ["freeze integrity", "GP age gate", "48-slice horizon", "future GP error", "Radar/RIPE blind holdout", "M-Lab RTT/throughput"],
        [
            "冻结完整性 12 个核心依赖文件通过；依赖审计 31 文件 aggregate hash 已记录。",
            "当前轨道年龄门禁、1584 星规模、约 43°/490 km 群组与 48 片/4 h 窗口通过。",
            "M-Lab RTT 比例约 0.9012/0.9019，吞吐中位数比例约 1.0142，但 P95 比例约 0.3135，尾部仍偏差明显。",
            "当前总门禁 8/12 通过、4 pending：未来 GP1、Radar 未来段、RIPE 精确样本数等仍需时间自然到达。",
        ],
        "14B 已证明冻结流程、数据新鲜度和因果协议成立；只有待未来数据到达并完成盲测后，才能对该窗口的轨道外推和网络性能真实性作最终结论。",
        "reports/experiment14b-prospective-external-validation/EXPERIMENT_14B_PROGRESS.md",
        "一次新鲜 GP0 并不能单独证明外推准确；必须保留 GP1 未来验证。当前 4 项 pending 不应被写成通过。",
        formula=(r"A_{end}=\max_i(T_{end}-epoch_i)\leq A_{max}", "8-5", "冻结窗口结束时的最大轨道数据年龄门禁。"),
    )

    # Chapter 9 conclusions
    add_heading(doc, "第 9 章  已证明、条件成立与尚未证明的结论", level=1, page_break=True)
    add_heading(doc, "9.1 当前已经有较强证据支持", level=2)
    add_table(
        doc,
        ["结论", "证据", "可使用的论文表述"],
        [
            ["模型能稳定生成动态真值", "构建/验收、三规模 profile、实验 1", "在公开轨道与研究级物理约束下生成可解释、可复现状态"],
            ["INT/OAM 非全知闭环已实现", "过程包、实验 7、前端与导出", "仅凭已交付 INT report 重构逐时间片全网状态"],
            ["增强方案改善多指标重构", "实验 2、10", "在中大型星座的等预算/相近预算设置下形成条件优势"],
            ["补全后端结论经多种子强基线检验", "实验 3：2 星座×3 mask seed×4 指标", "相同观测和字节下，后端效果具有明确的星座与指标依赖性"],
            ["动态 LEO 需要拓扑适配", "实验 8", "静态 reference plan 在受控动态压力下显著失效"],
            ["local repair 减少规划工作量", "实验 12", "Telesat 三个 48 片自然窗口中降低确定性规划工作量且质量非劣"],
            ["选择性 metadata 具备逐包意义", "实验 13", "小规模 ns-3 中减少 header/MTU 压力并改善交付/AoI"],
            ["外部验证协议已冻结", "实验 14B", "新鲜 GP0、有限窗口、未来 GP1 不回调的可复现协议"],
        ],
        widths=[2650, 2600, 4110],
        font_size=8.4,
        first_col_bold=True,
    )
    add_heading(doc, "9.2 只能在条件范围内成立", level=2)
    add_bullets(doc, [
        "增强 LEO-INT-MC 优于原生：主要在中大型动态星座和预注册预算/窗口下成立，小型存在退化与边界。",
        "拓扑复用有贡献：当前正式证据是 Telesat 0% 额外压力下的 local repair，不是 exact reuse，也不能外推所有压力。",
        "通信开销降低：小/中规模实验 2 成立；Starlink 部分结果是质量提升但字节上升。规划计算下降和网络字节下降必须分开。",
        "外部真实性：公开轨道、用户 RTT、宏观业务和 M-Lab 部分趋势有支撑；CPU、电量、队列仍是物理公式驱动潜变量。",
        "聚合模型有效：正常压力下 ns-3 趋势较一致；极端过载队列 P95 误差很大。",
        "强补全基线结果对三组路径观测 mask 成立，但尚未扩展到三个独立轨道历元或独立业务窗口。",
    ])
    add_heading(doc, "9.3 当前未被证明", level=2)
    add_bullets(doc, [
        "增强方法在所有星座、所有指标和所有开销上全面优于原生 INT-MC。",
        "多指标联合 CP 张量补全稳定优于逐指标二维补全，且收益独立于物理投影。",
        "exact topology reuse 在中大型星座、多压力、多 seed 下具有稳定统计收益。",
        "当前模型复刻真实 Starlink 内部 ISL、路由、CPU、电量、队列和 INT 数据。",
        "Radar 0.9134 或 0.955 是完全独立盲测泛化；前者含校准，后者存在研究者暴露边界。",
        "完整 P4/Tofino 可部署；当前只有算法级 INT 与 ns-3 包级交叉验证。",
        "ST-GNN、CoSTCo 或 graph-regularized 在所有规模普遍最优。",
    ])
    # Chapter 10
    add_heading(doc, "第 10 章  当前成熟度、论文立足点与后续工作", level=1, page_break=True)
    add_callout(doc, "科学表述原则", "该项目最稳健的论文姿态不是宣称“高仿真且全面最优”，而是明确：公开可验证部分用外部数据检验，内部不可公开状态用物理模型生成；算法优势在严格时间因果、硬预算和指定动态 LEO 条件下报告，并公开适用边界与负结果。", "info")
    add_figure(doc, create_evidence_score_figure(), "图 8  当前项目成熟度与证据边界评估", width=Inches(6.25))
    add_heading(doc, "10.1 项目立足点是否稳健", level=2)
    add_paragraph(doc, "当前立足点已经比较稳健。研究问题真实且重要：动态 LEO 中全量 INT 开销高、静态 probe 计划易失效、Ground OAM 又必须重构节点与链路全网状态。项目不是简单复现 INT-MC，而是把它作为一个低开销思想基线，提出面向轨道可预测动态拓扑的主动测量与重构框架。三阶段系统、无 oracle 边界、硬预算、动态压力、强基线和 ns-3 交叉验证共同构成了比单一算法脚本更完整的证据链。")
    add_heading(doc, "10.2 核心创新如何写成论文", level=2)
    add_numbered(doc, [
        "核心创新一：拓扑版本化、风险感知、单位字节信息增益驱动的 LEO 主动遥测规划。在同一硬预算下联合选择 reuse/repair/fresh、probe 路径和逐跳 metadata。",
        "核心创新二：无 oracle 的 Ground OAM 闭环。使用 AoI、预测方差、模型分歧、报告冲突和拓扑风险决定下一片观测，不读取真值误差。",
        "核心创新三：面向节点—链路异构状态的结构/物理约束重构。active mask、轨道邻居、观测锁定、电量传播和保守耦合共同提高多指标一致性。",
        "系统贡献：提供大规模聚合仿真与小规模 ns-3 逐包交叉验证的组合实验框架，并开放逐时间片真值、INT 过程、OAM 结果和误差报告。",
    ])
    add_heading(doc, "10.3 面向 INFOCOM 仍需补强的关键证据", level=2)
    add_table(
        doc,
        ["优先级", "工作", "最经济实现", "预期补强"],
        [
            ["P0", "完成实验 14B 的 4 个 pending", "等待未来 GP1/Radar/RIPE 到达，只跑冻结评估，不回调", "外部真实性由协议通过变成真正前瞻盲测结果"],
            ["P0", "拓扑复用跨压力统计", "保留已有 18 自然场景，追加少量按相似度分层的配对规划回放", "区分 trigger rate 与 conditional gain，扩大 local repair 结论"],
            ["P1", "强补全基线跨历元复核", "当前三 mask seed 已完成；后续冻结少量独立轨道/业务窗口，只重放 completion", "把观测选择稳健性扩展为跨窗口分布稳健性"],
            ["P1", "统一算法伪代码与复杂度", "从现有 planner 输出正式 Algorithm 1 和 O(|A|log|A|) 近似分析", "减少“规则拼盘”印象"],
            ["P1", "Git 冻结与复现包", "提交当前工作树、创建 tag、清理可再生产物、固定容器/WSL 版本", "审稿复现和 artifact 可信度"],
            ["P2", "BMv2/P4 小原型", "4–8 跳 selective metadata 与 report 生成", "进一步支撑实现可行性，非当前绝对必需"],
        ],
        widths=[900, 2500, 3400, 2560],
        font_size=8.2,
        first_col_bold=True,
    )
    add_heading(doc, "10.4 当前代码健康与交付风险", level=2)
    add_bullets(doc, [
        "最近验收中 npm run build、experiment10/12/13/14B 相关测试和仪表盘审计均通过。Vite 主包约 2.69 MB，仍可继续做按页面代码拆分以改善首屏加载。",
        "实验 14B 当前 8/12 pass、4 pending，不应在文稿中把 pending 改写为完成。",
        "当前 Git 工作树存在较多 modified/untracked 文件；正式冻结前必须提交或明确排除，避免报告结果无法对应 commit。",
        "部分旧 Markdown/CSV 在 Windows PowerShell 显示为乱码，但 HTML 与前端已修复；本 DOCX 使用 UTF-8 源和 CJK 字体生成。",
        "大型实验在 16 GB 内存机器上主要受 Node.js 单线程评分、对象分配与磁盘 I/O 限制；采用逐场景独立进程、流式时间片和有界并发，不应盲目全量并发。",
    ])

    # Appendices
    add_heading(doc, "附录 A  关键符号与状态对象", level=1, page_break=True)
    add_table(
        doc,
        ["符号", "含义", "对象/单位"],
        [
            ["Gt=(Vt,Et)", "时间片 t 的活动卫星拓扑", "节点/链路集合"],
            ["νt", "拓扑版本：类别、哈希、变化边、下一变化时间", "规划上下文"],
            ["Ωt", "本时间片直接观测对象集合", "节点/链路 metadata"],
            ["Ut(o)", "对象状态不确定性", "归一化分数"],
            ["Rt(o)", "未来断链/极区/回传/切换风险", "概率或归一化分数"],
            ["Bt", "时间片硬遥测字节预算", "byte"],
            ["AoI", "距离上次直接观测的状态年龄", "slice/time"],
            ["XV/XE", "节点/链路多指标状态", "矩阵或张量"],
            ["activeMask", "当前对象是否物理存在/可用", "Boolean"],
            ["observedMask", "该位置是否由已交付 INT 直接观测", "Boolean"],
        ],
        widths=[1700, 5160, 2500],
        font_size=8.8,
        first_col_bold=True,
    )

    add_heading(doc, "附录 B  当前关键配置摘要", level=1, page_break=True)
    add_table(
        doc,
        ["类别", "参数", "当前研究配置/说明"],
        [
            ["能量", "太阳翼面积/效率", "2.0 m² / 0.28"],
            ["能量", "电池/充放电效率/最低 SoC", "1200 Wh / 0.95 / 20%"],
            ["功耗", "base/comm/compute/payload", "80/100/50/100 W"],
            ["路由", "算法", "congestion-aware-shortest-path"],
            ["队列", "carryover/capacity/timeout", "0.65 / 16384 MB / 30000 ms"],
            ["ISL", "基础配置", "4 方向天线；激光等效频率；2500 Mbps 研究参数"],
            ["SGL", "基础配置", "Ka；最低仰角 25°；report capacity 600 Mbps"],
            ["星座", "小/中/大", "Iridium 66 / Telesat 351 / Starlink 1584"],
            ["时间", "主实验", "48 时间片；具体步长和时域由 manifest 冻结"],
        ],
        widths=[1400, 3000, 4960],
        font_size=8.8,
        first_col_bold=True,
    )
    add_callout(doc, "参数解释", "表中参数用于可复现研究仿真和同模型相对比较。除公开轨道与明确引用的外部统计外，不应把这些数值写成 Starlink/Iridium/Telesat 运营商真实硬件规格。", "warning")

    add_heading(doc, "附录 C  代码与实验产物索引", level=1, page_break=True)
    code_rows = [
        ["第一阶段轨道/拓扑", "src/simulation/walker.ts; tle.ts; realTleCatalog.ts"],
        ["天线与链路预算", "src/simulation/antenna.ts; linkBudget.ts; spaceEnvironment.ts"],
        ["业务与导出", "src/simulation/traffic.ts; export.ts; types.ts"],
        ["INT probe 执行", "stage2-int/tools/probe-path-planner.mjs; probe-int-runner.mjs"],
        ["Ground OAM", "stage2-int/tools/ground-oam-reconstructor.mjs"],
        ["增强路径选择", "stage2-int/tools/int-mc-path-selector.mjs"],
        ["拓扑预测/统一目标", "predict-contact-plan.mjs; topology-versioned-*.mjs"],
        ["矩阵/张量重构", "int-mc-reconstructor.mjs; int-mc-joint-tensor-completion.mjs"],
        ["强补全基线", "int-mc-additional-completion-backends.mjs"],
        ["开销核算", "telemetry-byte-budget.mjs; scale-adaptive-telemetry-budget.mjs"],
        ["逐包验证", "stage3-system-validation/ns3/scratch/leo-int-system-validation.cc"],
        ["使用与数据字典", "README.md; USER_GUIDE.md; DATASET_SCHEMA_GUIDE.md"],
        ["正式实验", "reports/experiment1-* 至 reports/experiment14b-*"],
    ]
    add_table(doc, ["主题", "主要位置"], code_rows, widths=[2600, 6760], font_size=8.8, first_col_bold=True)

    add_heading(doc, "附录 D  推荐复现顺序", level=1, page_break=True)
    add_code(doc, """
npm install
npm run build
npm run verify:constellations
npm run verify:goal
npm run int:experiment -- --tasks examples/datasets/stage1-standard-traffic.csv --orbit tle-sgp4 --mode operational --algorithm path-balance
npm run test:experiment7
npm run test:experiment10
npm run test:experiment12
npm run test:experiment13
npm run test:experiment14b
npm run experiment14b:audit
""")
    add_paragraph(doc, "大型正式实验应优先复用第一阶段真值、候选路径和 pass-1 OAM，只重跑所需规划或补全阶段；每个场景使用独立进程并及时释放内存。具体命令以 package.json 和各实验目录 README/REPORT 为准。")

    add_heading(doc, "结语", level=1)
    add_paragraph(doc, "这几天的工作把项目从一个可视化 Walker 网络逐步推进为完整研究平台：第一阶段提供可配置、物理一致、业务驱动的卫星网络真值；第二阶段以非全知 INT 逐片感知并由 Ground OAM 重构；增强方法针对 LEO 的拓扑动态、规模、星地回传和节点—链路异构状态进行适配；实验从内部一致性扩展到多种子、消融、动态压力、强补全基线、拓扑复用、ns-3 逐包和前瞻外部冻结验证。项目已经具备论文初稿的技术骨架和主要证据，但最可信的写法仍是精确限定主张、保留负结果，并完成实验 14B pending、跨压力复用统计和中大型强基线多 seed。")
    add_callout(doc, "最终判断", "项目已经大致完成从模型、遥测、算法到系统验证的闭环，立足点稳健、创新点可成立。当前最重要的工作不再是继续增加大量功能，而是冻结代码、补齐少数关键外部/统计证据，并把统一算法和适用边界写清楚。", "success")

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_report()
    document.save(OUTPUT)
    print(f"Generated: {OUTPUT}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")


if __name__ == "__main__":
    main()
